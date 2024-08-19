from PyQt6.QtWidgets import *
from PyQt6.QtCore import *
from PyQt6.QtGui import *
from pathlib import Path
from modules.atas.regex_termo_homolog import *
from modules.atas.regex_sicaf import *
from modules.atas.processar_homologacao import ProgressDialog
from modules.atas.processar_sicaf import SICAFDialog
from modules.atas.relatorio_indicadores import RelatorioIndicadores
from modules.atas.utils import create_button, load_icons, apply_standard_style, limpar_quebras_de_linha
from modules.atas.data_utils import DatabaseDialog, PDFProcessingThread, atualizar_modelo_com_dados, save_to_dataframe, load_file_path, obter_arquivos_txt, ler_arquivos_txt
from modules.atas.canvas_gerar_atas import criar_pastas_com_subpastas, abrir_pasta, gerar_soma_valor_homologado, inserir_relacao_empresa, inserir_relacao_itens, adicione_texto_formatado
from modules.atas.streamlit_dialog import StreamlitDialog
from diretorios import *
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
# import contextily as ctx
import traceback
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm 
from datetime import datetime
import json
# import seaborn as sns
from modules.planejamento.utilidades_planejamento import DatabaseManager
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import streamlit as st
import subprocess

NUMERO_ATA_GLOBAL = None
GERADOR_NUMERO_ATA = None

class CustomTreeView(QTreeView):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.cnpj = ""

    def contextMenuEvent(self, event):
        index = self.indexAt(event.pos())
        if not index.isValid():
            return

        item = self.model().itemFromIndex(index)
        if item and " - " in item.text():
            # Usar QTextDocument para extrair texto puro a partir do HTML
            doc = QTextDocument()
            doc.setHtml(item.text())
            plain_text = doc.toPlainText()

            # Extração do CNPJ sem HTML
            if " - " in plain_text:
                self.cnpj = plain_text.split(' - ')[0]

            menu = QMenu(self)
            copyAction = QAction(f"Copiar CNPJ {self.cnpj}", self)
            copyAction.triggered.connect(lambda: self.copy_cnpj(plain_text))
            menu.addAction(copyAction)
            menu.exec(event.globalPos())

    def copy_cnpj(self, text):
        if " - " in text:
            self.cnpj = text.split(' - ')[0]  # Extrai o CNPJ
        QApplication.clipboard().setText(self.cnpj)
        QToolTip.showText(QCursor.pos(), f"CNPJ {self.cnpj} copiado para área de transferência", self)

    def mousePressEvent(self, event):
        index = self.indexAt(event.pos())
        if index.isValid() and event.button() == Qt.MouseButton.LeftButton:
            # Expande ou colapsa o item clicado
            self.setExpanded(index, not self.isExpanded(index))
            
            # Se o item foi expandido, expanda também o primeiro nível de subitens
            if self.isExpanded(index):
                model = self.model()
                numRows = model.rowCount(index)
                for row in range(numRows):
                    childIndex = model.index(row, 0, index)
                    self.setExpanded(childIndex, True)
                    # Colapsa todos os subníveis abaixo do primeiro nível
                    self.collapseAllChildren(childIndex)

        super().mousePressEvent(event)

    def collapseAllChildren(self, parentIndex):
        """Recursivamente colapsa todos os subníveis de um dado índice."""
        model = self.model()
        numRows = model.rowCount(parentIndex)
        for row in range(numRows):
            childIndex = model.index(row, 0, parentIndex)
            self.collapseAllChildren(childIndex)
            self.setExpanded(childIndex, False)
class GerarAtasWidget(QWidget):
    def __init__(self, icons_dir, parent=None):
        super().__init__(parent)
        self.icons_dir = Path(icons_dir)
        self.buttons = {}
        self.tr_variavel_df_carregado = None 
        # Diretórios configurados
        self.pdf_dir = Path(PDF_DIR)
        self.txt_dir = Path(TXT_DIR)
        self.sicaf_dir = Path(SICAF_DIR)
        self.sicaf_txt_dir = Path(SICAF_TXT_DIR)
            
        self.mapeamento_colunas = self.obter_mapeamento_colunas()
        self.current_dataframe = None
        self.pe_pattern = None
        self.setup_ui()
        self.progressDialog = ProgressDialog(self.pdf_dir, self)
        self.setup_pdf_processing_thread()        
        self.db_manager = DatabaseManager(CONTROLE_DADOS)

    def verificar_ou_criar_diretorios(self, diretorios):
        for diretorio in diretorios:
            if not diretorio.exists():
                diretorio.mkdir(parents=True, exist_ok=True)

    def obter_mapeamento_colunas(self):
        return {
            "Grupo": "grupo",
            "Item": "item_num",
            "Catálogo": "catalogo",
            "Descrição": "descricao_tr",
            "Descrição Detalhada": "descricao_detalhada",
            "Unidade": "unidade",
            "Quantidade": "quantidade",
            "Valor Estimado": "valor_estimado",
            "Valor Homologado": "valor_homologado_item_unitario",
            "Desconto (%)": "percentual_desconto",
            "Valor Estimado Total": "valor_estimado_total_do_item",
            "Valor Homologado Total": "valor_homologado_total_item",
            "Marca Fabricante": "marca_fabricante",
            "Modelo Versão": "modelo_versao",
            "UASG": "uasg",
            "Órgão Responsável": "orgao_responsavel",
            "Número": "num_pregao",
            "Ano": "ano_pregao",
            "SRP": "srp",
            "Objeto": "objeto",
            "Situação": "situacao",
            "Melhor Lance": "melhor_lance",
            "Valor Negociado": "valor_negociado",
            "Ordenador Despesa": "ordenador_despesa",
            "Empresa": "empresa",
            "CNPJ": "cnpj",
            "Endereço": "endereco",
            "CEP": "cep",
            "Município": "municipio",
            "Telefone": "telefone",
            "Email": "email",
            "Responsável Legal": "responsavel_legal"
        }
    
    def setup_ui(self):
        self.main_layout = QVBoxLayout(self)

        # Configurar o alerta com tamanho mínimo
        self.setup_alert_label()

        # Configurar os botões
        self.setup_buttons_up()

        # Adicionar um widget vazio que atuará como um espaçador no topo
        spacer_top = QWidget()
        spacer_top.setSizePolicy(QSizePolicy.Policy.Minimum, QSizePolicy.Policy.Minimum)
        self.main_layout.addWidget(spacer_top)

        # Adicionar a mensagem centralizada e definir uma área fixa para ela
        self.message_label = QLabel("Simplicidade é o último degrau da sabedoria\n(Khalil Gibran)")
        self.message_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.message_label.setStyleSheet("font-size: 16pt; font-style: italic; padding: 20px;")
        
        # Definir um tamanho fixo para a área onde o message_label ou treeView será exibido
        self.fixed_area_widget = QWidget()
        self.fixed_area_layout = QVBoxLayout(self.fixed_area_widget)
        self.fixed_area_layout.addWidget(self.message_label, alignment=Qt.AlignmentFlag.AlignCenter)
        
        # Garantir que o fixed_area_widget ocupe o máximo de espaço disponível
        self.main_layout.addWidget(self.fixed_area_widget)
        self.fixed_area_widget.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)

        # Adicionar um widget vazio que atuará como um espaçador na parte inferior
        spacer_bottom = QWidget()
        spacer_bottom.setSizePolicy(QSizePolicy.Policy.Minimum, QSizePolicy.Policy.Minimum)
        self.main_layout.addWidget(spacer_bottom)

        # Setup da parte inferior dos botões
        self.setup_buttons_down()

        self.setLayout(self.main_layout)
        self.setMinimumSize(1000, 580)

    def setup_treeview(self):
        # Remover ou ocultar a mensagem quando o treeView for exibido
        if self.message_label:
            self.message_label.hide()

        self.model = QStandardItemModel()  # Inicializando o modelo se ainda não foi inicializado
        self.treeView = CustomTreeView()
        self.treeView.setModel(self.model)
        
        # Remover a mensagem e adicionar o treeView na área fixa
        self.fixed_area_layout.removeWidget(self.message_label)
        self.message_label.deleteLater()  # Opcional: deletar o label para liberar memória
        self.fixed_area_layout.addWidget(self.treeView)

        self.treeView.header().setDefaultAlignment(Qt.AlignmentFlag.AlignCenter)
        self.treeView.setAnimated(True)  # Facilita a visualização da expansão/colapso
        self.treeView.setUniformRowHeights(True)  # Uniformiza a altura das linhas      
        self.treeView.setItemsExpandable(True)  # Garantir que o botão para expandir esteja visível
        self.treeView.setExpandsOnDoubleClick(False)  # Evita a expansão por duplo clique
        self.setup_treeview_styles()
        
    def setup_alert_label(self):
        icon_path = str(self.icons_dir / 'alert.png')
        text = (f"<img src='{icon_path}' style='vertical-align: middle;' width='24' height='24'> "
                "Pressione '<b><u>Termo de Referência</u></b>' para adicionar os dados 'Catálogo', "
                "'Descrição' e 'Descrição Detalhada' do Termo de Referência. "
                f"<img src='{icon_path}' style='vertical-align: middle;' width='24' height='24'>")
        self.alert_label = QLabel(text)
        self.alert_label.setStyleSheet("font-size: 12pt; padding: 5px;")
        self.alert_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.main_layout.addWidget(self.alert_label)
        self.atasDialog = None

    def setup_buttons_generic(self, button_definitions):
        buttons_layout = QHBoxLayout()
        self.icons = load_icons(self.icons_dir)
        for name, icon_key, callback, tooltip, animate in button_definitions:
            icon = self.icons.get(icon_key, None)
            button = create_button(name, icon, callback, tooltip, QSize(30, 30), QSize(120, 30), None)
            self.buttons[name] = button
            buttons_layout.addWidget(button)
        return buttons_layout

    def setup_buttons_up(self):
        button_definitions = self.obter_definicoes_botoes()
        self.main_layout.addLayout(self.setup_buttons_generic(button_definitions))

    def setup_buttons_down(self):
        button_definitions = self.obter_definicoes_botoes_embaixo()
        self.main_layout.addLayout(self.setup_buttons_generic(button_definitions))

    def obter_definicoes_botoes(self):
        return [
            ("Tabela Vazia", 'add-task', self.tabela_vazia, "Importe um arquivo .xlsx com 4 colunas com índice 'item_num', 'catalogo', 'descricao_tr' e 'descricao_detalada'.", True),
            ("Termo de Referência", 'priority', self.import_tr, "Importe um arquivo .xlsx com 4 colunas com índice 'item_num', 'catalogo', 'descricao_tr' e 'descricao_detalada'.", True),
            ("Termo de Homologação", 'data-collection', self.processar_homologacao, "Faça o download dos termos de homologação e mova para a pasta de processamento dos Termos de Homologação", False),
            ("SICAF", 'sicaf', self.processar_sicaf, "Faça o download do SICAF (Nível I - Credenciamento) e mova para a pasta de processamento do SICAF", False),
        ]

    def obter_definicoes_botoes_embaixo(self):
        return [
            ("Ata / Contrato", 'verify_menu', self.abrir_dialog_atas, "Com o database concluíodo é possível gerar as atas ou contratos", False),
            ("Database", 'data-processing', self.update_database, "Salva ou Carrega os dados do Database", False),
            ("Salvar Tabela", 'table', self.salvar_tabela, "Importe um arquivo .xlsx com 4 colunas com índice 'item_num', 'catalogo', 'descricao_tr' e 'descricao_detalada'.", True),
            ("Indicadores", 'dashboard', self.dashboard_indicadores, "Visualize os indicadores do relatório", False),
        ]

    def tabela_vazia(self):
        # Definindo o nome do arquivo XLSX
        arquivo_xlsx = str(self.sicaf_dir / "tabela_vazia.xlsx")

        # Definindo as colunas da tabela
        colunas = [
            "item_num", "catalogo", "descricao_tr", "descricao_detalhada", 
        ]

        # Criando um DataFrame com a coluna item_num numerada até a linha 10
        df_vazio = pd.DataFrame({
            "item_num": range(1, 11),
            "catalogo": [""] * 10,
            "descricao_tr": [""] * 10,
            "descricao_detalhada": [""] * 10
        })

        try:
            # Tentar abrir o arquivo em modo exclusivo para escrita para verificar se ele está em uso
            with open(arquivo_xlsx, 'w') as f:
                pass  # O arquivo pode ser aberto e escrito

            # Salvando o DataFrame como um arquivo XLSX
            df_vazio.to_excel(arquivo_xlsx, index=False)

            # Ajustando a largura das colunas usando openpyxl
            workbook = load_workbook(arquivo_xlsx)
            worksheet = workbook.active

            worksheet.column_dimensions['A'].width = 15  # item_num
            worksheet.column_dimensions['B'].width = 100  # catalogo
            worksheet.column_dimensions['C'].width = 150  # descricao_tr
            worksheet.column_dimensions['D'].width = 200  # descricao_detalhada

            workbook.save(arquivo_xlsx)

            # Abrindo o arquivo XLSX gerado
            os.startfile(arquivo_xlsx)

        except PermissionError:
            # Se houver um PermissionError, significa que o arquivo está aberto em outra aplicação
            QMessageBox.warning(self, "Arquivo Aberto", "A tabela 'tabela_vazia.xlsx' está aberta. Feche o arquivo antes de tentar salvá-lo novamente.")
        
    def setup_treeview_styles(self):
        header = self.treeView.header()
        header.setSectionResizeMode(QHeaderView.ResizeMode.ResizeToContents)

    def setup_pdf_processing_thread(self):
        if not self.pdf_dir.exists() or not self.txt_dir.exists():
            QMessageBox.critical(self, "Erro", "Diretório de PDF ou TXT não encontrado.")
            return
        self.processing_thread = PDFProcessingThread(self.pdf_dir, self.txt_dir)
        self.processing_thread.progress_updated.connect(self.progressDialog.update_progress)
        self.processing_thread.processing_complete.connect(self.progressDialog.on_conversion_finished)

    def import_tr(self):
        # Oculta o message_label ao importar
        if self.message_label:
            self.message_label.hide()

        # Abrir um diálogo de arquivo para selecionar o arquivo
        file_path, _ = QFileDialog.getOpenFileName(self, "Importar Termo de Referência", "", 
                                                    "Arquivos Excel (*.xlsx);;Arquivos LibreOffice (*.ods)")
        if not file_path:
            return  # Se o usuário cancelar, não faça nada

        # Verifica a extensão do arquivo para usar o método adequado
        ext = Path(file_path).suffix.lower()

        if ext == '.xlsx':
            # Para arquivos Excel (.xlsx)
            df = pd.read_excel(file_path)
        elif ext == '.ods':
            # Para arquivos LibreOffice (.ods)
            df = pd.read_excel(file_path, engine='odf')
        else:
            QMessageBox.warning(self, "Formato não suportado", "Formato de arquivo não suportado.")
            return

        # Formatar e validar os dados (se necessário)
        erros = self.formatar_e_validar_dados(df)
        if erros:
            QMessageBox.critical(self, "Erro de Validação", "\n".join(erros))
            return

        # Atualizar o DataFrame carregado
        self.tr_variavel_df_carregado = df

        # Verifique se o treeView já foi inicializado, se não, inicialize-o
        if not hasattr(self, 'treeView'):
            self.setup_treeview()

        # Atualizar a visualização da tabela com os dados importados
        self.atualizar_modelo_com_dados(df)

        # Atualizar o alerta para informar o usuário que o Termo de Referência foi carregado
        self.atualizar_alerta_apos_importar_tr()

        # Exibir uma mensagem informando que o arquivo foi importado com sucesso
        QMessageBox.information(self, "Importação Concluída", f"Arquivo {Path(file_path).name} importado com sucesso.")

    def formatar_e_validar_dados(self, df):
        erros = []

        # Verificar se as colunas obrigatórias estão presentes
        colunas_obrigatorias = ["item_num", "catalogo", "descricao_tr", "descricao_detalhada"]
        for coluna in colunas_obrigatorias:
            if coluna not in df.columns:
                erros.append(f"A coluna obrigatória '{coluna}' está faltando no arquivo.")

        # Se houver erros nas colunas obrigatórias, retornar imediatamente
        if erros:
            return erros

        # Remover quebras de linha e espaços extras
        for col in colunas_obrigatorias:
            df[col] = df[col].apply(lambda x: " ".join(str(x).replace("\n", " ").split()) if pd.notnull(x) else x)

        # Verificar se item_num é inteiro e sequencial
        df['item_num'] = pd.to_numeric(df['item_num'], errors='coerce')

        if df['item_num'].isnull().any():
            erros.append("A coluna 'item_num' contém valores não numéricos.")

        # Verificar sequencialidade
        sequencia = df['item_num'].dropna().astype(int).sort_values().unique()
        esperado = list(range(sequencia.min(), sequencia.max() + 1))
        if not all(item in sequencia for item in esperado):
            faltando = set(esperado) - set(sequencia)
            erros.append(f"Números sequenciais faltando em 'item_num': {sorted(faltando)}")

        # Verificar se catalogo, descricao_tr, e descricao_detalhada são textos
        df['catalogo'] = df['catalogo'].astype(str)
        df['descricao_tr'] = df['descricao_tr'].astype(str)
        df['descricao_detalhada'] = df['descricao_detalhada'].astype(str)

        # Verificar valores em branco nas colunas obrigatórias
        for coluna in colunas_obrigatorias:
            if df[coluna].isnull().any() or df[coluna].eq("").any():
                erros.append(f"A coluna '{coluna}' contém valores em branco.")

        # Verificar se há erros
        if erros:
            erros.insert(0, "Não foi possível carregar a tabela devido aos seguintes erros:")

        return erros

    def atualizar_modelo_com_dados(self, df_relevante):
        self.model.clear()
        self.model.setHorizontalHeaderLabels(['Item', 'Catálogo', 'Descrição', 'Descrição Detalhada'])
        for _, row in df_relevante.iterrows():
            item_num = QStandardItem(str(row['item_num']))
            catalogo = QStandardItem(str(row['catalogo']))
            descricao_tr = QStandardItem(str(row['descricao_tr']))
            descricao_detalhada = QStandardItem(str(row['descricao_detalhada']))
            item_num.setEditable(False)
            catalogo.setEditable(False)
            descricao_tr.setEditable(False)
            descricao_detalhada.setEditable(False)
            self.model.appendRow([item_num, catalogo, descricao_tr, descricao_detalhada])
        self.treeView.expandAll()
        for column in range(self.model.columnCount()):
            self.treeView.resizeColumnToContents(column)

    def atualizar_alerta_apos_importar_tr(self):
        icon_path = str(self.icons_dir / 'confirm.png')
        new_text = (f"<img src='{icon_path}' style='vertical-align: middle;' width='24' height='24'> "
                    "Salve os Termos de Homologação na pasta correta e pressione '<b><u>Termo de Homologação</u></b>' para processar os dados. "
                    f"<img src='{icon_path}' style='vertical-align: middle;' width='24' height='24'>")
        self.alert_label.setText(new_text)

    def atualizar_alerta_apos_processar_homologacao(self):
        icon_path = str(self.icons_dir / 'sicaf.png')
        new_text = (f"<img src='{icon_path}' style='vertical-align: middle;' width='24' height='24'> "
                    "Clique com o botão direito no TreeView para copiar o CNPJ para facilitar a busca do SICAF Nível I. "
                    f"<img src='{icon_path}' style='vertical-align: middle;' width='24' height='24'>")
        self.alert_label.setText(new_text)

    def processar_homologacao(self):
        # Verifica se a pasta de PDFs existe
        if not self.pdf_dir.exists():
            QMessageBox.warning(self, "Erro", "Pasta de PDFs não encontrada.")
            return
        
        # Lista todos os arquivos PDF na pasta
        pdf_files = list(self.pdf_dir.glob("*.pdf"))
        
        # Inicializa o ProgressDialog e processa todos os arquivos, sem verificar por novos arquivos
        total_files = len(pdf_files)
        self.progressDialog = ProgressDialog(total_files, self.pdf_dir, self)
        self.progressDialog.processing_complete.connect(self.finalizar_processamento_homologacao)
        self.progressDialog.show()

    def finalizar_processamento_homologacao(self, extracted_data):
        self.homologacao_dataframe = save_to_dataframe(extracted_data, self.tr_variavel_df_carregado, DATABASE_DIR, self.current_dataframe)
        
        if self.homologacao_dataframe is not None:
            self.current_dataframe = self.homologacao_dataframe  # Atualiza o DataFrame corrente
            self.update_treeview_with_dataframe(self.homologacao_dataframe)
            self.atualizar_alerta_apos_processar_homologacao()
            
            # Verifica se as colunas necessárias existem
            if {'num_pregao', 'ano_pregao', 'uasg'}.issubset(self.current_dataframe.columns):
                # Filtra linhas onde qualquer uma das colunas chave contém NaN
                filtered_df = self.current_dataframe.dropna(subset=['num_pregao', 'ano_pregao', 'uasg'])
                
                # Gera o nome da tabela apenas para linhas sem NaN
                def create_table_name(row):
                    return f"{row['num_pregao']}-{row['ano_pregao']}-{row['uasg']}-Homolog"

                filtered_df['table_name'] = filtered_df.apply(create_table_name, axis=1)
                
                # Debugging output
                print("Valores de 'num_pregao':", filtered_df['num_pregao'].unique())
                print("Valores de 'ano_pregao':", filtered_df['ano_pregao'].unique())
                print("Valores de 'uasg':", filtered_df['uasg'].unique())
                print("Nomes de tabelas gerados:", filtered_df['table_name'].unique())
                
                if filtered_df['table_name'].nunique() == 1:
                    table_name = filtered_df['table_name'].iloc[0]
                    self.save_data(table_name)  # Chama a função de salvar com o nome da tabela
                else:
                    QMessageBox.critical(self, "Erro", "A combinação de 'num_pregao', 'ano_pregao', e 'uasg' não é única. Por favor, verifique os dados.")
            else:
                QMessageBox.critical(self, "Erro", "Dados necessários para criar o nome da tabela não estão presentes.")
            return self.current_dataframe  # Retorna o DataFrame atualizado
        else:
            QMessageBox.warning(self, "Erro", "Falha ao salvar os dados.")
            return None  # Retorna None para indicar que o processo falhou

    def save_data(self, table_name):
        if isinstance(self.current_dataframe, pd.DataFrame) and not self.current_dataframe.empty:
            print("Salvando DataFrame com as colunas:", self.current_dataframe.columns)
            try:
                with self.db_manager as conn:
                    self.current_dataframe.to_sql(table_name, conn, if_exists='replace', index=False)
                QMessageBox.information(self, "Sucesso", f"DataFrame salvo com sucesso em '{table_name}'!")
            except Exception as e:
                QMessageBox.critical(self, "Erro", f"Erro ao salvar os dados no banco de dados: {e}")
        else:
            QMessageBox.critical(self, "Erro", "Nenhum DataFrame válido disponível para salvar ou o objeto não é um DataFrame.")

    def processar_sicaf(self):
        if self.current_dataframe is not None:
            dataframe_to_use = self.current_dataframe
        else:
            QMessageBox.warning(self, "Erro", "Primeiro processe a homologação ou carregue os dados do banco de dados.")
            return

        if not self.sicaf_dir.exists():
            QMessageBox.warning(self, "Erro", "Pasta de PDFs não encontrada.")
            return

        self.progressSicafDialog = SICAFDialog(self.sicaf_dir, dataframe_to_use, self)
        # Conecta o sinal a ambos os métodos
        self.progressSicafDialog.processing_complete.connect(self.finalizar_processamento_sicaf)
        self.progressSicafDialog.processing_complete.connect(self.receber_df_final)
        self.progressSicafDialog.show()

    def receber_df_final(self, df_final):
        if isinstance(df_final, pd.DataFrame):
            self.current_dataframe = df_final  # Atualize o DataFrame atual
            print("DataFrame final recebido do SICAF:")
            print(df_final)

            # Verifica se as colunas necessárias existem
            if {'num_pregao', 'ano_pregao', 'uasg'}.issubset(self.current_dataframe.columns):
                # Filtra linhas onde qualquer uma das colunas chave contém NaN
                filtered_df = self.current_dataframe.dropna(subset=['num_pregao', 'ano_pregao', 'uasg'])
                
                # Gera o nome da tabela apenas para linhas sem NaN
                def create_table_name(row):
                    return f"{row['num_pregao']}-{row['ano_pregao']}-{row['uasg']}-Homolog-Sicaf"

                filtered_df['table_name'] = filtered_df.apply(create_table_name, axis=1)
                
                # Debugging output
                print("Valores de 'num_pregao':", filtered_df['num_pregao'].unique())
                print("Valores de 'ano_pregao':", filtered_df['ano_pregao'].unique())
                print("Valores de 'uasg':", filtered_df['uasg'].unique())
                print("Nomes de tabelas gerados:", filtered_df['table_name'].unique())
                
                if filtered_df['table_name'].nunique() == 1:
                    table_name = filtered_df['table_name'].iloc[0]
                    self.save_data(table_name)  # Chama a função de salvar com o nome da tabela
                else:
                    QMessageBox.critical(self, "Erro", "A combinação de 'num_pregao', 'ano_pregao', e 'uasg' não é única. Por favor, verifique os dados.")
            else:
                QMessageBox.critical(self, "Erro", "Dados necessários para criar o nome da tabela não estão presentes.")
            return self.current_dataframe  # Retorna o DataFrame atualizado
        else:
            QMessageBox.warning(self, "Erro", "Dados recebidos não são válidos.")

    def finalizar_processamento_sicaf(self, extracted_data):
        if isinstance(extracted_data, pd.DataFrame):
            print("DataFrame resultante do SICAF:")
            print(extracted_data)
            self.update_treeview_with_dataframe(extracted_data)
        else:
            print("Erro: Dados recebidos não são um DataFrame.")
            QMessageBox.warning(self, "Erro", "Os dados recebidos não são válidos.")

    def handle_loaded_data(self, loaded_dataframe, pe_pattern=None):
        if isinstance(loaded_dataframe, pd.DataFrame) and not loaded_dataframe.empty:
            self.current_dataframe = loaded_dataframe  # Atualiza o DataFrame corrente
            self.pe_pattern = pe_pattern  # Armazena o padrão PE identificado
            print(f"DataFrame atualizado e carregado:\n{self.current_dataframe.head()}")
            print(f"Padrão PE identificado: {self.pe_pattern}")

            # Verifique se o treeView já foi inicializado, se não, inicialize-o
            if not hasattr(self, 'treeView'):
                self.setup_treeview()

            self.update_treeview_with_dataframe(self.current_dataframe)
        else:
            QMessageBox.warning(self, "Aviso", "Os dados carregados não são um DataFrame válido ou estão vazios.")

    def update_treeview_with_dataframe(self, dataframe):
        if dataframe is None:
            QMessageBox.critical(self, "Erro", "O DataFrame não está disponível para atualizar a visualização.")
            return
        
        creator = ModeloTreeview(self.icons_dir)
        self.model = creator.criar_modelo(dataframe)
        
        # Verifique se o treeView está inicializado
        if not hasattr(self, 'treeView'):
            self.setup_treeview()

        self.treeView.setModel(self.model)
        # self.treeView.setItemDelegate(HTMLDelegate())
        self.treeView.reset()


    def update_database(self):
        # Sempre abre o diálogo, independentemente da existência de um DataFrame atual
        dialog = DatabaseDialog(self, self.current_dataframe, self.handle_loaded_data)
        dialog.exec()

    def update_progress(self, value):
        if self.progressDialog.isVisible():
            self.progressDialog.progressBar.setValue(value)
        else:
            # Caso a barra de progresso não esteja visível, você pode optar por mostrá-la aqui
            self.progressDialog.show()
            self.progressDialog.progressBar.setValue(value)
                    
    def abrir_dialog_atas(self):
        if self.current_dataframe is not None:
            dataframe_to_use = self.current_dataframe
            if all(col in dataframe_to_use.columns for col in ['empresa', 'num_pregao', 'ano_pregao']):
                print("Colunas de 'empresa', 'num_pregao', 'ano_pregao' do DataFrame:")
                print(dataframe_to_use[['empresa', 'num_pregao', 'ano_pregao']])
            else:
                print("Alguma das colunas 'empresa', 'num_pregao', 'ano_pregao' não está presente no DataFrame.")
        else:
            dataframe_to_use = None
            print("Nenhum DataFrame atual disponível.")

        if self.atasDialog is None or not self.atasDialog.isVisible():
            self.atasDialog = AtasDialog(self, pe_pattern=self.pe_pattern, dataframe=dataframe_to_use)
            self.atasDialog.dataframe_updated.connect(self.on_dataframe_updated)  # Conectar ao método que trata o DataFrame
            self.atasDialog.show()
        else:
            self.atasDialog.raise_()
            self.atasDialog.activateWindow()

    def on_dataframe_updated(self, updated_dataframe):
        # Atualiza o DataFrame atual com as modificações recebidas
        self.current_dataframe = updated_dataframe
        print("DataFrame atualizado recebido do diálogo Atas:")

        # Verifica se o DataFrame é None ou se as colunas necessárias estão ausentes
        if self.current_dataframe is None or not {'numero_ata', 'item_num'}.issubset(self.current_dataframe.columns):
            return 

        print(self.current_dataframe[['numero_ata', 'item_num']])

        # Verifica se as colunas necessárias existem
        if {'num_pregao', 'ano_pregao', 'uasg'}.issubset(self.current_dataframe.columns):
            # Filtra linhas onde qualquer uma das colunas chave contém NaN
            filtered_df = self.current_dataframe.dropna(subset=['num_pregao', 'ano_pregao', 'uasg'])
            
            # Gera o nome da tabela apenas para linhas sem NaN
            def create_table_name(row):
                return f"{row['num_pregao']}-{row['ano_pregao']}-{row['uasg']}-Final"

            filtered_df['table_name'] = filtered_df.apply(create_table_name, axis=1)
            
            # Debugging output
            print("Valores de 'num_pregao':", filtered_df['num_pregao'].unique())
            print("Valores de 'ano_pregao':", filtered_df['ano_pregao'].unique())
            print("Valores de 'uasg':", filtered_df['uasg'].unique())
            print("Nomes de tabelas gerados:", filtered_df['table_name'].unique())
            
            if filtered_df['table_name'].nunique() == 1:
                table_name = filtered_df['table_name'].iloc[0]
                self.save_data(table_name)  # Chama a função de salvar com o nome da tabela
            else:
                QMessageBox.critical(self, "Erro", "A combinação de 'num_pregao', 'ano_pregao', e 'uasg' não é única. Por favor, verifique os dados.")
        else:
            QMessageBox.critical(self, "Erro", "Dados necessários para criar o nome da tabela não estão presentes.")

    def dashboard_indicadores(self):
        if self.current_dataframe is not None:
            # Salva o DataFrame em um arquivo CSV temporário
            self.current_dataframe.to_csv(STREAMLIT_CSV, index=False)
            
            if STREAMLIT_CSV.exists():
                # Configura a variável de ambiente para não solicitar email
                os.environ["EMAIL_OPT_OUT"] = "true"
                
                # Abrir o diálogo do Streamlit
                dialog = StreamlitDialog(self)
                dialog.exec()
            else:
                QMessageBox.warning(self, "Aviso", "Erro ao salvar o arquivo CSV.")
        else:
            QMessageBox.warning(self, "Aviso", "Não há dados para salvar.")


    def salvar_tabela(self):
        if self.current_dataframe is not None:
            # Define o caminho do arquivo a ser salvo
            arquivo_excel = str(self.pdf_dir / 'TabelaAtual.xlsx')
            # Salva o DataFrame no arquivo Excel
            self.current_dataframe.to_excel(arquivo_excel, index=False)
            # Abre o arquivo Excel
            os.startfile(arquivo_excel)
        else:
            QMessageBox.warning(self, "Aviso", "Não há dados para salvar.")

    def indicadores_normceim(self):
        if self.current_dataframe is not None:
            # Supondo que pe_pattern é armazenado em algum lugar após ser determinado
            self.dialogo_indicadores = RelatorioIndicadores(dataframe=self.current_dataframe, parent=self, pe_pattern=self.pe_pattern)
            self.dialogo_indicadores.show()
        else:
            QMessageBox.warning(self, "Aviso", "Não há dados carregados.")


class ModeloTreeview:
    def __init__(self, icons_dir):
        print(f"Carregando ícones de: {icons_dir}")
        self.icons = load_icons(icons_dir)

    def determinar_itens_iguais(self, row, empresa_items):
        empresa_name = str(row['empresa']) if pd.notna(row['empresa']) else ""
        cnpj = str(row['cnpj']) if pd.notna(row['cnpj']) else ""
        situacao = str(row['situacao']) if pd.notna(row['situacao']) else "Não definido"
        is_situacao_only = not empresa_name and not cnpj
        parent_key = f"{situacao}" if is_situacao_only else f"{cnpj} - {empresa_name}".strip(" -")
        
        if parent_key not in empresa_items:
            parent_item = QStandardItem()
            parent_item.setEditable(False)
            icon_key = 'alert' if is_situacao_only else ('checked' if pd.notna(row['endereco']) else 'unchecked')
            icon = self.icons.get(icon_key, QIcon())  # Obtém o ícone ou um ícone vazio se não encontrado
            parent_item.setIcon(icon)
            return parent_key, parent_item

        return parent_key, None

    def update_view(self, view):
        # Força a atualização da view para garantir que os ícones sejam exibidos
        view.viewport().update()
        
    def criar_modelo(self, dataframe):
        model, header = self.initializar_modelo(dataframe)
        empresa_items = self.processar_linhas(dataframe, model)
        self.atualizar_contador_cabecalho(empresa_items, model)
        return model

    def initializar_modelo(self, dataframe):
        total_items = len(dataframe)
        situacoes_analizadas = ['Adjudicado e Homologado', 'Fracassado e Homologado', 'Deserto e Homologado', 'Anulado e Homologado']
        nao_analisados = len(dataframe[~dataframe['situacao'].isin(situacoes_analizadas)])
        header = f"Total de itens da licitação {total_items} | Total de itens analisados {total_items - nao_analisados} | Total de itens não analisados {nao_analisados}"
        model = QStandardItemModel()
        model.setHorizontalHeaderLabels([header])
        return model, header

    def processar_linhas(self, dataframe, model):
        empresa_items = {}
        for _, row in dataframe.iterrows():
            self.processar_linhas_individualmente(row, model, empresa_items)
        return empresa_items

    def processar_linhas_individualmente(self, row, model, empresa_items):
        parent_key, parent_item = self.determinar_itens_iguais(row, empresa_items)
        if parent_item:
            model.appendRow(parent_item)
            if parent_key not in empresa_items:
                empresa_items[parent_key] = {
                    'item': parent_item,
                    'count': 0,
                    'details_added': False,
                    'items_container': QStandardItem()  # Não defina o texto aqui
                }
                empresa_items[parent_key]['items_container'].setEditable(False)

        # Ajuste para incrementar a contagem em todas as situações
        if parent_key in empresa_items:
            empresa_items[parent_key]['count'] += 1

        self.adicionar_informacao_ao_item(row, empresa_items[parent_key]['item'], empresa_items, parent_key)

    def adicionar_informacao_ao_item(self, row, parent_item, empresa_items, parent_key):
        font_size = "14px"  # Define o tamanho da fonte
        situacoes_especificas = ['Fracassado e Homologado', 'Deserto e Homologado', 'Anulado e Homologado']
        situacao = row['situacao']

        # Determinar se a situação é específica ou "Não definido"
        if situacao not in situacoes_especificas and situacao != 'Adjudicado e Homologado':
            situacao = 'Não definido'

        if situacao in situacoes_especificas or situacao == 'Não definido':
            # Cria um item com informações básicas sem detalhes extras para situações específicas
            item_text = f"Item {row['item_num']} - {row['descricao_tr']} - {situacao}"
            item_info = QStandardItem(item_text)
            item_info.setEditable(False)
            parent_item.appendRow(item_info)
        else:
            # Processo normal para 'Adjudicado e Homologado'
            if not empresa_items[parent_key]['details_added']:
                self.adicionar_detalhes_empresa(row, parent_item)
                empresa_items[parent_key]['items_container'].setText("")  # Limpa o texto se necessário
                parent_item.appendRow(empresa_items[parent_key]['items_container'])
                empresa_items[parent_key]['details_added'] = True

            # Adicionando itens específicos da licitação
            self.adicionar_subitens_detalhados(row, empresa_items[parent_key]['items_container'])

        # Atualizar o texto do container com base na contagem de itens
        item_count_text = "Item" if empresa_items[parent_key]['count'] == 1 else "Relação de itens:"
        empresa_items[parent_key]['items_container'].setText(f"{item_count_text} ({empresa_items[parent_key]['count']})")
        
    def atualizar_contador_cabecalho(self, empresa_items, model):
        font_size = "16px"  # Definir o tamanho da fonte para os cabeçalhos dos itens
        for chave_item_pai, empresa in empresa_items.items():
            count = empresa['count']
            # Formatar o texto com HTML para ajustar o tamanho da fonte
            display_text = f"{chave_item_pai} (1 item)" if count == 1 else f"{chave_item_pai} ({count} itens)"
            empresa['item'].setText(display_text)

    def adicionar_detalhes_empresa(self, row, parent_item):
        infos = [
            f"Endereço: {row['endereco']}, CEP: {row['cep']}, Município: {row['municipio']}" if pd.notna(row['endereco']) else "Endereço: Não informado",
            f"Contato: {row['telefone']} Email: {row['email']}" if pd.notna(row['telefone']) else "Contato: Não informado",
            f"Responsável Legal: {row['responsavel_legal']}" if pd.notna(row['responsavel_legal']) else "Responsável Legal: Não informado"
        ]
        for info in infos:
            info_item = QStandardItem(info)
            info_item.setEditable(False)
            parent_item.appendRow(info_item)

    def criar_dados_sicaf_do_item(self, row):
        fields = ['endereco', 'cep', 'municipio', 'telefone', 'email', 'responsavel_legal']
        return [self.criar_detalhe_item(field.capitalize(), row[field]) for field in fields if pd.notna(row[field])]

    def adicionar_subitens_detalhados(self, row, sub_items_layout):
        item_info_html = f"Item {row['item_num']} - {row['descricao_tr']} - {row['situacao']}"
        item_info = QStandardItem(item_info_html)
        item_info.setEditable(False)
        sub_items_layout.appendRow(item_info)

        detalhes = [
            f"Descrição Detalhada: {row['descricao_detalhada']}",
            f"Unidade de Fornecimento: {row['unidade']} Quantidade: {self.formatar_quantidade(row['quantidade'])} "
            f"Valor Estimado: {self.formatar_brl(row['valor_estimado'])} Valor Homologado: {self.formatar_brl(row['valor_homologado_item_unitario'])} "
            f"Desconto: {self.formatar_percentual(row['percentual_desconto'])} Marca: {row['marca_fabricante']} Modelo: {row['modelo_versao']}",
        ]

        for detalhe in detalhes:
            detalhe_item = QStandardItem(detalhe)
            detalhe_item.setEditable(False)
            item_info.appendRow(detalhe_item)


    def criar_detalhe_item(self, label, data):
        return QStandardItem(f"<b>{label}:</b> {data if pd.notna(data) else 'Não informado'}")

    def formatar_brl(self, valor):
        try:
            if valor is None:
                return "Não disponível"  # ou outra representação adequada para seu caso de uso
            return f"R$ {float(valor):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        except ValueError:
            return "Valor inválido"

    def formatar_quantidade(self, valor):
        try:
            float_value = float(valor)
            if float_value.is_integer():
                return f"{int(float_value)}"
            else:
                return f"{float_value:.2f}".replace('.', ',')
        except ValueError:
            return "Erro de Formatação"

    def formatar_percentual(self, valor):
        try:
            percent_value = float(valor)
            return f"{percent_value:.2f}%"
        except ValueError:
            return "Erro de Formatação"
        
class AtasDialog(QDialog):
    NUMERO_ATA_GLOBAL = None  # Defina isso em algum lugar adequado dentro de sua classe
    dataframe_updated = pyqtSignal(object)  # Sinal para emitir o DataFrame atualizado
    pastas_criadas = set()  # Rastreador de pastas criadas

    def __init__(self, parent=None, pe_pattern=None, dataframe=None):
        super().__init__(parent)
        self.db_manager = DatabaseManager(CONTROLE_DADOS)
        self.pe_pattern = pe_pattern
        self.nup_data = None
        self.dataframe = dataframe 
        self.settings = QSettings("YourCompany", "YourApp")  # Ajuste esses valores para seu aplicativo
        self.configurar_ui()

    def closeEvent(self, event):
        # Quando o diálogo é fechado, emite o DataFrame atualizado
        self.dataframe_updated.emit(self.dataframe)
        super().closeEvent(event)

    def update_title_label(self):
        html_text = (
            f"<span style='font-size: 28px'>Painel de Geração de Atas/Contratos</span>"
        )
        self.titleLabel.setText(html_text)
        self.titleLabel.setTextFormat(Qt.TextFormat.RichText)
        self.titleLabel.setStyleSheet(" font-size: 32px; font-weight: bold;")

        header_layout = QHBoxLayout()
        header_layout.addWidget(self.titleLabel)
        header_layout.addSpacerItem(QSpacerItem(40, 20, QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum))
        # Adicionando os botões antes do spacer e do pixmap
        self.add_action_buttons(header_layout)
        pixmap = QPixmap(str(MARINHA_PATH))
        pixmap = pixmap.scaled(80, 80, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
        image_label = QLabel()
        image_label.setPixmap(pixmap)
        header_layout.addWidget(image_label)

        return header_layout

    def add_action_buttons(self, layout):
        # Caminhos para os ícones
        icon_x = QIcon(str(ICONS_DIR / "cancel.png"))  # Caminho para o ícone de cancelamento
        
        # Criação dos botões
        button_x = self.create_button("  Sair", icon_x, self.reject, "Cancelar alterações e fechar", QSize(130, 50), QSize(30, 30))
        
        # Adicionando os botões ao layout
        layout.addWidget(button_x)
        self.apply_widget_style(button_x)

    def apply_widget_style(self, widget):
        widget.setStyleSheet("font-size: 14pt;") 

    def create_button(self, text, icon=None, callback=None, tooltip_text="", button_size=None, icon_size=None):
        btn = QPushButton(text, self)
        if icon:
            btn.setIcon(icon)
            btn.setIconSize(icon_size if icon_size else QSize(40, 40))
        if button_size:
            btn.setFixedSize(button_size)
        btn.setToolTip(tooltip_text)
        btn.setFont(QFont('Arial', 14))
        if callback:
            btn.clicked.connect(callback)
        return btn

    def salvar_alteracoes(self):
        texto_header = self.header_editor.toHtml()  # Captura o texto HTML do QTextEdit
        salvar_configuracoes({'header_text': texto_header})

    def configurar_ui(self):
        self.setWindowTitle("Atas / Contratos")
        self.setFont(QFont('Arial', 14))

        # Cria um QVBoxLayout principal para o QDialog
        layout = QVBoxLayout(self)
        self.resize(800, 650)

        # Cria e configura o titleLabel
        self.titleLabel = QLabel()
        header_layout = self.update_title_label()
        layout.addLayout(header_layout)

        header_label = QLabel("Editar Cabeçalho:")
        header_label.setFont(QFont('Arial', 14))
        layout.addWidget(header_label)
        
        self.header_editor = QTextEdit()
        self.header_editor.setFont(QFont('Arial', 12))
        self.header_editor.setMinimumHeight(180)

        configuracoes = carregar_configuracoes()
        initial_text = configuracoes.get('header_text', '')  # Carrega o texto salvo ou usa string vazia
        self.header_editor.setHtml(initial_text)
        layout.addWidget(self.header_editor)
                
        # # Texto inicial com HTML para formatação
        # initial_text = ("A União, por intermédio do CENTRO DE INTENDÊNCIA DA MARINHA EM BRASÍLIA (CeIMBra), com sede na "
        #                 "Esplanada dos Ministérios, Bloco “N”, Prédio Anexo, 2º andar, CEP: 70055-900, na cidade de Brasília – DF, "
        #                 "inscrito(a) sob o CNPJ nº 00.394.502/0594-67, neste ato representado pelo Capitão de Fragata (IM) "
        #                 "Thiago Martins Amorim, Ordenador de Despesa, nomeado(a) pela Portaria nº 241 de 25 de abril de 2024, "
        #                 "do Com7°DN, c/c Ordem de Serviço nº 57/2024 de 25 de abril de 2024 do CeIMBra, considerando o "
        #                 "julgamento da licitação na modalidade de pregão, na forma eletrônica, para REGISTRO DE PREÇOS nº "
        #                 "<span style='color: blue;'>{{num_pregao}}</span>/2024, processo administrativo nº <span style='color: blue;'>{{nup}}</span>, RESOLVE registrar os preços da(s) "
        #                 "empresa(s) indicada(s) e qualificada(s) nesta ATA, de acordo com a classificação por ela(s) alcançada(s) "
        #                 "e na(s) quantidade(s) cotada(s), atendendo as condições previstas no Edital de licitação, sujeitando-se "
        #                 "as partes às normas constantes na Lei nº 14.133, de 1º de abril de 2021, no Decreto n.º 11.462, de "
        #                 "31 de março de 2023, e em conformidade com as disposições a seguir:")
        # self.header_editor.setHtml(initial_text)
        # layout.addWidget(self.header_editor)
        layout.addItem(QSpacerItem(20, 10, QSizePolicy.Policy.Minimum, QSizePolicy.Policy.Fixed))

        # Configurar o combobox das cidades
        cidades_label = QLabel("Selecione Cidade:")
        cidades_label.setFont(QFont('Arial', 14))
        cidades_label.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Fixed)
        layout.addWidget(cidades_label)

        self.cidades_combobox = QComboBox()
        cidades = ["Brasília-DF", "Rio Grande-RS", "Salvador-BA", "São Pedro da Aldeia-RJ", "Rio de Janeiro-RJ", "Natal-RN", "Manaus-MA"]
        self.cidades_combobox.addItems(cidades)
        self.cidades_combobox.setFont(QFont('Arial', 14))
        self.cidades_combobox.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
        self.cidades_combobox.setMaximumWidth(300)  # Ajusta a largura máxima do combobox

        # Carregar a última cidade selecionada
        last_city = self.settings.value("last_selected_city", "Brasília-DF")
        index = self.cidades_combobox.findText(last_city)
        if index != -1:
            self.cidades_combobox.setCurrentIndex(index)
        
        layout.addWidget(self.cidades_combobox)

        layout.addItem(QSpacerItem(20, 10, QSizePolicy.Policy.Minimum, QSizePolicy.Policy.Fixed))

        # Configuração para seleção da organização gerenciadora
        org_label = QLabel("Selecione a Organização Gerenciadora:")
        org_label.setFont(QFont('Arial', 14))
        layout.addWidget(org_label)

        self.org_combobox = QComboBox()
        organizations = [
            "Centro de Intendência da Marinha em Brasília (CeIMBra)",
            "Centro de Intendência da Marinha em Natal (CeIMNa)",
            "Centro de Intendência da Marinha em Manaus (CeIMMa)",
            "Centro de Intendência da Marinha em Salvador (CeIMSa)",
            "Centro de Intendência da Marinha em Rio Grande (CeIMRG)",
            "Centro de Intendência da Marinha em São Pedro da Aldeia (CeIMSPA)",
            "Hospital Naval de Brasília (HNBra)",
            "Hospital Naval Marcílio Dias (HNMD)",
        ]
        self.org_combobox.addItems(organizations)
        self.org_combobox.setFont(QFont('Arial', 14))
        self.org_combobox.setMaximumWidth(700)
        last_org = self.settings.value("last_selected_org", organizations[0])
        self.org_combobox.setCurrentText(last_org)
        layout.addWidget(self.org_combobox)

        layout.addItem(QSpacerItem(20, 10, QSizePolicy.Policy.Minimum, QSizePolicy.Policy.Fixed))

        self.configurar_rotulos(layout)
        # self.configurar_entrada_e_botao_confirmar(layout)
        # self.configurar_botoes_acao(layout)
        self.carregar_e_exibir_ultimo_contrato()
        self.criar_botao_gerar_ata(layout)
        self.setLayout(layout)

    def closeEvent(self, event):
        # Salvar as últimas seleções
        self.settings.setValue("last_selected_city", self.cidades_combobox.currentText())
        self.settings.setValue("last_selected_org", self.org_combobox.currentText())
        self.dataframe_updated.emit(self.dataframe)
        super().closeEvent(event)

    def configurar_rotulos(self, layout):
        self.rotulo_ultimo_contrato = QLabel("O último contrato gerado foi:")
        self.rotulo_ultimo_contrato.setFont(QFont('Arial', 14))
        layout.addWidget(self.rotulo_ultimo_contrato)

        rotulo = QLabel("Digite o próximo Número de Controle de Atas/Contratos:")
        rotulo.setFont(QFont('Arial', 14))

        # Criar um QHBoxLayout para o rótulo e o campo de entrada
        linha_rotulo_entrada = QHBoxLayout()
        linha_rotulo_entrada.addWidget(rotulo)

        self.entradaAta = QLineEdit(self)
        self.entradaAta.setFont(QFont('Arial', 14))
        self.entradaAta.setPlaceholderText("Digite um número até 4 dígitos")
        self.entradaAta.setMaxLength(10)
        self.entradaAta.setFixedWidth(self.entradaAta.fontMetrics().horizontalAdvance('0000') + 20)
        linha_rotulo_entrada.addWidget(self.entradaAta)

        # Usando a função create_button para criar o botão de confirmar
        self.botaoConfirmar = self.create_button("Confirmar", None, self.confirmar_numero_ata_e_nup_do_processo, "Clique para confirmar o número")
        linha_rotulo_entrada.addWidget(self.botaoConfirmar)

        # Adiciona o layout horizontal ao layout vertical principal
        layout.addLayout(linha_rotulo_entrada)
        
    def carregar_e_exibir_ultimo_contrato(self):
        ultimo_numero_contrato = self.carregar_ultimo_contrato()
        if ultimo_numero_contrato is not None:
            self.atualizar_rotulo_ultimo_contrato(f"Nº {ultimo_numero_contrato}")
            self.entradaAta.setText(str(ultimo_numero_contrato + 1))
        else:
            self.rotulo_ultimo_contrato.setText("O último número de ata/contrato gerado foi: Nenhum")

    def criar_botao_gerar_ata(self, layout):
        # Caminhos para os ícones
        icon_confirm = QIcon(str(ICONS_DIR / "production_red.png"))  # Caminho para o ícone de confirmação
        
        # Criação do botão
        button_confirm = self.create_button("  Gerar Atas de Registro de Preços", icon_confirm, self.gerar_ata_de_registro_de_precos, "Após inserir as informações, clique para gerar as Atas", QSize(380, 50), QSize(50, 50))
        
        # Criação de um layout horizontal para centralizar o botão
        hbox = QHBoxLayout()
        hbox.addSpacerItem(QSpacerItem(40, 20, QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum))
        hbox.addWidget(button_confirm)
        hbox.addSpacerItem(QSpacerItem(40, 20, QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum))
        
        # Adicionando o layout horizontal ao layout principal
        layout.addLayout(hbox)
        
        # Aplicar estilo ao botão
        self.apply_widget_style(button_confirm)

    @staticmethod
    def convert_pe_format(pe_string):
        if pe_string is None:
            print("Erro: pe_string é None")
            return None

        pe_formatted = pe_string.replace('PE-', 'PE ').replace('-', '/')
        print(f"Converted PE format: {pe_formatted}")  # Depuração
        return pe_formatted

    def obter_nup(self, pe_formatted):
        try:
            with self.db_manager as conn:
                query = f"SELECT nup FROM controle_processos WHERE id_processo LIKE '%{pe_formatted}%'"
                df = pd.read_sql(query, conn)
                if not df.empty:
                    self.nup_data = {
                        'nup': df.iloc[0]['nup']
                    }
                    return self.nup_data
                else:
                    return None
        except Exception as e:
            print(f"Erro ao acessar o banco de dados: {e}")
            return None
                                       
    def atualizar_rotulo_ultimo_contrato(self, ultimo_numero_contrato):
        self.rotulo_ultimo_contrato.setText(f"O último número de ata/contrato gerado foi: {ultimo_numero_contrato}")

    def salvar_ultimo_contrato(self, ultimo_num_contrato):
        with open(ULTIMO_CONTRATO_DIR, "w") as f:
            f.write(str(ultimo_num_contrato))  # Convertendo para string

    def carregar_ultimo_contrato(self):
        try:
            with open(ULTIMO_CONTRATO_DIR, "r") as f:
                return int(f.read().split('/')[-1])
        except (FileNotFoundError, ValueError):
            return None

    def confirmar_numero_ata_e_nup_do_processo(self):
        numero_ata = self.entradaAta.text()
        print(f"Valor inserido pelo usuário: '{numero_ata}'")  # Use aspas para capturar espaços em branco ou strings vazias

        if numero_ata.isdigit() and len(numero_ata) <= 4:
            AtasDialog.NUMERO_ATA_GLOBAL = int(numero_ata)
            
            # Verificação para self.pe_pattern
            if self.pe_pattern is None:
                QMessageBox.warning(self, "Erro", "Padrão de pregão não encontrado. Por favor, carregue um database antes de continuar.")
                print("Padrão de pregão não encontrado. Necessário carregar um database.")
                return
            
            pe_formatted = self.convert_pe_format(self.pe_pattern)
            nup = self.obter_nup(pe_formatted)

            if nup:
                self.nup_data = nup
                QMessageBox.information(self, "Número Confirmado", f"Número da ata definido para: {numero_ata}")
                print(f"Número da ATA confirmado e definido como {numero_ata}.")
            else:
                QMessageBox.warning(self, "Erro ao Obter NUP", "Não foi encontrado um padrão de pregão para prosseguir. Por favor, carregue um database antes de continuar.")
                print("Não foi encontrado um padrão de pregão para prosseguir. Necessário carregar um database.")
        else:
            QMessageBox.warning(self, "Número Inválido", "Por favor, digite um número válido de até 4 dígitos.")
            print("Tentativa de inserção de um número inválido.")


    def gerar_ata_de_registro_de_precos(self):
        if not self.nup_data:  # Verifica se nup_data está vazia ou é None
            self.nup_data = "(INSIRA O NUP)"  # Atribui um valor padrão caso não exista nup_data
        
        # Chama salvar_alteracoes para garantir que todas as mudanças sejam salvas antes de processar
        self.salvar_alteracoes()
        
        self.processar_ata_de_registro_de_precos(self.nup_data, self.dataframe)

    def processar_ata_de_registro_de_precos(self, nup_data, dataframe):
        if AtasDialog.NUMERO_ATA_GLOBAL is None:
            QMessageBox.information(self, "Inserir Número da ATA", "Por favor, insira o número da ATA para continuar.")
            return

        criar_pastas_com_subpastas(dataframe)  # Chamando função externa de criação de pastas
        ultimo_num_ata = self.processar_ata(AtasDialog.NUMERO_ATA_GLOBAL, nup_data, dataframe)

        self.salvar_ultimo_contrato(ultimo_num_ata)
        self.atualizar_rotulo_ultimo_contrato(ultimo_num_ata)

    def processar_ata(self, NUMERO_ATA, nup_data, dataframe):
        if isinstance(nup_data, dict) and 'nup' in nup_data:
            nup = nup_data['nup']
        else:
            nup = "(INSIRA O NUP)" 
        relatorio_path = get_relatorio_path()
        combinacoes = dataframe[['uasg', 'num_pregao', 'ano_pregao', 'empresa']].drop_duplicates().values
        NUMERO_ATA_atualizado = NUMERO_ATA

        # Verifica se a coluna 'numero_ata' existe, senão, cria ela
        if 'numero_ata' not in dataframe.columns:
            dataframe['numero_ata'] = None

        for uasg, num_pregao, ano_pregao, empresa in combinacoes:
            if not pd.isna(num_pregao) and not pd.isna(ano_pregao) and not pd.isna(empresa):
                path_dir_principal, path_subpasta = self.preparar_diretorios(relatorio_path, num_pregao, ano_pregao, empresa)
                registros_empresa = dataframe[dataframe['empresa'] == empresa]
                NUMERO_ATA_atualizado, num_contrato = self.processar_empresa(registros_empresa, empresa, path_subpasta, nup, NUMERO_ATA_atualizado)

                # Atualizar o DataFrame com o número de contrato atualizado para a empresa processada
                dataframe.loc[dataframe['empresa'] == empresa, 'numero_ata'] = num_contrato

        abrir_pasta(str(path_dir_principal))
        print(dataframe[['numero_ata', 'item_num']])  # Mostra os valores das colunas 'numero_ata' e 'item_num'
        return NUMERO_ATA_atualizado

    def preparar_diretorios(self, relatorio_path, num_pregao, ano_pregao, empresa):        
        nome_dir_principal = f"PE {int(num_pregao)}-{int(ano_pregao)}"
        path_dir_principal = relatorio_path / nome_dir_principal
        nome_empresa_limpa = self.limpar_nome_empresa(empresa)
        path_subpasta = path_dir_principal / nome_empresa_limpa

        # Verifica se a pasta já foi criada anteriormente
        chave_pasta = (nome_dir_principal, nome_empresa_limpa)
        if chave_pasta not in AtasDialog.pastas_criadas:
            if not path_subpasta.exists():
                path_subpasta.mkdir(parents=True, exist_ok=True)
                print(f"Criado subdiretório: {path_subpasta}")
            AtasDialog.pastas_criadas.add(chave_pasta)
        else:
            print(f"Subdiretório já existente, não será recriado: {path_subpasta}")

        return path_dir_principal, path_subpasta

    def limpar_nome_empresa(self, nome_empresa):
        # Substituir '/' e ':' por sublinhado
        caracteres_para_sublinhado = ['/', ':']
        for char in caracteres_para_sublinhado:
            nome_empresa = nome_empresa.replace(char, '_')
        
        # Substituir '.' por nada (remover)
        nome_empresa = nome_empresa.replace('.', '')

        # Substituir outros caracteres inválidos por sublinhados
        caracteres_invalidos = ['<', '>', '_', '"', '\\', '|', '?', '*']
        for char in caracteres_invalidos:
            nome_empresa = nome_empresa.replace(char, '_')

        # Remover espaços extras e sublinhados no final do nome da empresa
        nome_empresa = nome_empresa.rstrip(' _')

        # Substituir múltiplos espaços ou sublinhados consecutivos por um único sublinhado
        nome_empresa = '_'.join(filter(None, nome_empresa.split(' ')))

        return nome_empresa
    
    def processar_empresa(self, registros_empresa, empresa, path_subpasta, nup, NUMERO_ATA_atualizado):
        if not registros_empresa.empty:
            registro = registros_empresa.iloc[0].to_dict()
            itens_relacionados = registros_empresa.to_dict('records')
            context, num_contrato = self.criar_contexto(registro, empresa, NUMERO_ATA_atualizado, nup, itens_relacionados)
            
            # Passando num_contrato para salvar_documento
            self.salvar_documento(path_subpasta, empresa, context, registro, itens_relacionados, num_contrato)

            # Garanta que NUMERO_ATA_atualizado é um inteiro antes de incrementar
            NUMERO_ATA_atualizado = int(NUMERO_ATA_atualizado) + 1  # Atualiza o número da ATA após processar com sucesso
        else:
            print(f"Nenhum registro encontrado para a empresa: {empresa}")
            num_contrato = None
        return NUMERO_ATA_atualizado, num_contrato


    def criar_contexto(self, registro, empresa, NUMERO_ATA_atualizado, nup, itens_relacionados):
        ano_atual = datetime.now().year
        num_contrato = f"{registro['uasg']}/{ano_atual}-{NUMERO_ATA_atualizado:03}/00"
        texto_substituto = f"Pregão Eletrônico nº {registro['num_pregao']}/{registro['ano_pregao']}\n{num_contrato}"
        soma_valor_homologado = gerar_soma_valor_homologado(itens_relacionados)

        dados_da_ug_contratante_cabecalho = self.header_editor.toPlainText()
        cidade_selecionada = self.cidades_combobox.currentText()
        organizacao_selecionada = self.org_combobox.currentText()

        return ({
            "num_pregao": registro['num_pregao'],
            "ano_pregao": registro['ano_pregao'],
            "empresa": empresa,
            "uasg": registro['uasg'],
            "numero_ata": NUMERO_ATA_atualizado,
            "soma_valor_homologado": soma_valor_homologado,
            "cabecalho": texto_substituto,
            "dados_ug_contratante": dados_da_ug_contratante_cabecalho,
            "contrato": num_contrato,
            "endereco": registro["endereco"],
            "cnpj": registro["cnpj"],
            "objeto": registro["objeto"],
            "ordenador_despesa": registro["ordenador_despesa"],
            "responsavel_legal": registro["responsavel_legal"],
            "nup": nup,
            "email": registro["email"],
            "cidade": cidade_selecionada,
            "organizacao": organizacao_selecionada
        }, num_contrato)

    def salvar_email(self, path_subpasta, context):
        nome_arquivo_txt = "E-mail.txt"
        path_arquivo_txt = path_subpasta / nome_arquivo_txt
        with open(path_arquivo_txt, "w") as arquivo_txt:
            texto_email = (f"{context['email']}\n\n"
                        f"Sr. Representante.\n\n"
                        f"Encaminho em anexo a Vossa Senhoria a ATA {context['contrato']} "
                        f"decorrente do Pregão Eletrônico (SRP) nº {context['num_pregao']}/{context['ano_pregao']}, do Centro "
                        f"de Intendência da Marinha em Brasília (CeIMBra).\n\n"
                        f"Os documentos deverão ser conferidos, assinados e devolvidos a este Comando.\n\n"
                        f"A empresa receberá uma via, devidamente assinada, após a publicação.\n\n"
                        f"Respeitosamente,\n")
            arquivo_txt.write(texto_email)

    def salvar_documento(self, path_subpasta, empresa, context, registro, itens_relacionados, num_contrato):
        max_len = 40  # Definindo o limite máximo para o nome da empresa
        contrato_limpo = self.limpar_nome_empresa(num_contrato)[:max_len].rstrip()

        # Preparar o template do documento
        tpl = DocxTemplate(TEMPLATE_PATH)
        tpl.render(context)

        # Montar o nome do arquivo, garantindo que não ultrapasse os limites comuns de sistemas de arquivos
        nome_documento = f"{contrato_limpo}.docx"
        path_documento = path_subpasta / nome_documento

        # Salvar o documento
        tpl.save(path_documento)

        # Alterar o documento após a criação inicial para incluir informações detalhadas
        self.alterar_documento_criado(path_documento, registro, registro["cnpj"], itens_relacionados)

        # Salvando o arquivo de email associado
        self.salvar_email(path_subpasta, context)

    def alterar_documento_criado(self, caminho_documento, registro, cnpj, itens):
        # Carregar o documento
        doc = Document(caminho_documento)

        # Iterar por cada parágrafo
        for paragraph in doc.paragraphs:
            if '{relacao_empresa}' in paragraph.text:
                paragraph.clear()
                inserir_relacao_empresa(paragraph, registro, cnpj)

            if '{relacao_item}' in paragraph.text:
                paragraph.clear()
                inserir_relacao_itens(paragraph, itens)

        # Extrair o diretório do caminho do documento
        diretorio_documento = os.path.dirname(caminho_documento)

        # Gerar o arquivo Excel no mesmo diretório do documento
        caminho_arquivo_excel = os.path.join(diretorio_documento, 'relacao_itens.xlsx')
        gerar_excel_relacao_itens(itens, caminho_arquivo_excel)

        # Inserir a tabela do Excel no local do marcador <<tabela_itens>>
        inserir_tabela_do_excel_no_documento(doc, caminho_arquivo_excel)

        # Salvar o documento atualizado
        doc.save(caminho_documento)

    def inserir_relacao_empresa(self, paragrafo, registro, cnpj):
        dados = {
            "Razão Social": registro["empresa"],
            "CNPJ": registro["cnpj"],
            "Endereço": registro["endereco"],
            "Município-UF": registro["municipio"],
            "CEP": registro["cep"],
            "Telefone": registro["telefone"],
            "E-mail": registro["email"]
        }

        total_itens = len(dados)
        contador = 1
        
        for chave, valor in dados.items():
            adicione_texto_formatado(paragrafo, f"{chave}: ", True)

            # Verifica se é a penúltima linha
            if contador == total_itens - 1:
                adicione_texto_formatado(paragrafo, f"{valor}; e\n", False)
            # Verifica se é a última linha
            elif contador == total_itens:
                adicione_texto_formatado(paragrafo, f"{valor}.\n", False)
            else:
                adicione_texto_formatado(paragrafo, f"{valor};\n", False)

            contador += 1
        
        adicione_texto_formatado(paragrafo, "Representada neste ato, por seu representante legal, o(a) Sr(a) ", False)
        adicione_texto_formatado(paragrafo, f'{registro["responsavel_legal"]}.\n', False)

from openpyxl import load_workbook
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def inserir_tabela_do_excel_no_documento(doc, caminho_arquivo_excel):
    # Carregar a planilha do Excel
    wb = load_workbook(caminho_arquivo_excel)
    ws = wb.active

    for paragraph in doc.paragraphs:
        if '<<tabela_itens>>' in paragraph.text:
            # Remover o texto do marcador
            paragraph.text = paragraph.text.replace('<<tabela_itens>>', '')

            # Criar uma nova tabela no documento Word
            table = doc.add_table(rows=0, cols=3)

            for i in range(0, ws.max_row, 4):
                # Primeira linha (mesclada e pintada de cinza claro)
                row_cells = table.add_row().cells
                row_cells[0].merge(row_cells[2])
                run = row_cells[0].paragraphs[0].add_run(str(ws.cell(row=i+1, column=1).value))
                run.font.size = Pt(12)
                run.font.bold = True
                shading_elm = parse_xml(r'<w:shd {} w:fill="D3D3D3"/>'.format(nsdecls('w')))
                row_cells[0]._element.get_or_add_tcPr().append(shading_elm)

                # Segunda linha (mesclada com "Descrição Detalhada:" em negrito e quebras de linha)
                row_cells = table.add_row().cells
                row_cells[0].merge(row_cells[2])
                
                # Adiciona quebra de linha antes de "Descrição Detalhada:"
                run = row_cells[0].paragraphs[0].add_run("\n")
                
                # Adiciona "Descrição Detalhada:" em negrito
                run = row_cells[0].paragraphs[0].add_run("Descrição Detalhada:")
                run.font.size = Pt(10)
                run.font.bold = True

                # Adicionando o texto restante após "Descrição Detalhada:"
                texto_segunda_linha = str(ws.cell(row=i+2, column=1).value)
                run = row_cells[0].paragraphs[0].add_run(f" {texto_segunda_linha}")
                run.font.size = Pt(10)
                
                # Adiciona quebra de linha após o texto
                row_cells[0].paragraphs[0].add_run("\n")

                # Terceira linha (manter formatação padrão)
                row_cells = table.add_row().cells
                for j in range(3):
                    value = ws.cell(row=i+3, column=j+1).value
                    if value is not None:
                        texto = str(value)
                        if j == 0 and texto.startswith("UF:"):
                            # Negrito apenas para "UF:"
                            run = row_cells[j].paragraphs[0].add_run("UF:")
                            run.font.size = Pt(10)
                            run.font.bold = True
                            
                            # Texto que segue "UF:"
                            run = row_cells[j].paragraphs[0].add_run(texto[3:])
                            run.font.size = Pt(10)
                        elif j == 1 and texto.startswith("Marca:"):
                            # Negrito apenas para "Marca:"
                            run = row_cells[j].paragraphs[0].add_run("Marca:")
                            run.font.size = Pt(10)
                            run.font.bold = True
                            
                            # Texto que segue "Marca:"
                            run = row_cells[j].paragraphs[0].add_run(texto[6:])
                            run.font.size = Pt(10)
                        elif j == 2 and texto.startswith("Modelo:"):
                            # Negrito apenas para "Modelo:"
                            run = row_cells[j].paragraphs[0].add_run("Modelo:")
                            run.font.size = Pt(10)
                            run.font.bold = True
                            
                            # Texto que segue "Modelo:"
                            run = row_cells[j].paragraphs[0].add_run(texto[7:])
                            run.font.size = Pt(10)
                        else:
                            row_cells[j].text = texto
                            row_cells[j].paragraphs[0].runs[0].font.size = Pt(10)

                # Quarta linha (manter formatação padrão)
                row_cells = table.add_row().cells
                for j in range(3):
                    value = ws.cell(row=i+4, column=j+1).value
                    if value is not None:
                        texto = str(value)
                        if j == 0 and texto.startswith("Quantidade:"):
                            # Negrito apenas para "Quantidade:"
                            run = row_cells[j].paragraphs[0].add_run("Quantidade:")
                            run.font.size = Pt(10)
                            run.font.bold = True

                            # Texto que segue "Quantidade:"
                            run = row_cells[j].paragraphs[0].add_run(texto[11:])
                            run.font.size = Pt(10)
                            row_cells[j].paragraphs[0].add_run("\n")  # Adiciona quebra de linha
                        elif j == 1 and texto.startswith("Valor Unitário:"):
                            # Negrito apenas para "Valor Unitário:"
                            run = row_cells[j].paragraphs[0].add_run("Valor Unitário:")
                            run.font.size = Pt(10)
                            run.font.bold = True

                            # Texto que segue "Valor Unitário:"
                            run = row_cells[j].paragraphs[0].add_run(texto[15:])
                            run.font.size = Pt(10)
                            row_cells[j].paragraphs[0].add_run("\n")  # Adiciona quebra de linha
                        elif j == 2 and texto.startswith("Valor Total:"):
                            # Negrito apenas para "Valor Total:"
                            run = row_cells[j].paragraphs[0].add_run("Valor Total:")
                            run.font.size = Pt(10)
                            run.font.bold = True

                            # Texto que segue "Valor Total:"
                            run = row_cells[j].paragraphs[0].add_run(texto[12:])
                            run.font.size = Pt(10)
                            row_cells[j].paragraphs[0].add_run("\n")  # Adiciona quebra de linha
                        else:
                            row_cells[j].text = texto
                            row_cells[j].paragraphs[0].runs[0].font.size = Pt(10)
                            row_cells[j].paragraphs[0].add_run("\n")  # Adiciona quebra de linha
            # Mover a tabela para logo após o parágrafo atual
            move_table_after_paragraph(paragraph, table)
            break

def move_table_after_paragraph(paragraph, table):
    # Move a tabela para ficar logo após o parágrafo atual
    tbl, p = table._tbl, paragraph._element
    p.addnext(tbl)


def salvar_configuracoes(dados):
    with open('configuracoes.json', 'w') as arquivo:
        json.dump(dados, arquivo)

def carregar_configuracoes():
    try:
        with open('configuracoes.json', 'r') as arquivo:
            return json.load(arquivo)
    except (FileNotFoundError, json.JSONDecodeError):
        return {}  # Retorna um dicionário vazio se o arquivo não existir ou estiver corrompido

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill

def format_currency(value):
    """ Função para formatar valores monetários no formato brasileiro. """
    print(f"Valor original: {value}")  # Print para depuração

    if isinstance(value, str):
        # Tentar converter string para float
        try:
            value = float(value.replace('.', '').replace(',', '.'))
        except ValueError:
            return value  # Se a conversão falhar, retorna o valor original
    elif not isinstance(value, (int, float)):
        return value  # Retorna o valor original se não for um número

    # Formatar o valor como moeda brasileira
    formatted_value = f"R$ {value:,.2f}".replace(',', 'temp').replace('.', ',').replace('temp', '.')
    
    print(f"Valor formatado: {formatted_value}")  # Print para depuração
    return formatted_value

def gerar_excel_relacao_itens(itens, caminho_arquivo_excel='relacao_itens.xlsx'):
    wb = Workbook()
    ws = wb.active

    fonte_tamanho_12_cinza = Font(size=12, bold=True)
    fundo_cinza = PatternFill(start_color="D3D3D3", end_color="D3D3D3", fill_type="solid")
    fonte_tamanho_10 = Font(size=10)

    linha_atual = 1

    for item in itens:
        ws.merge_cells(start_row=linha_atual, start_column=1, end_row=linha_atual, end_column=3)
        cell_1 = ws.cell(row=linha_atual, column=1, value=f"Item: {item['item_num']} - {item['descricao_tr']} ({item['catalogo']})")
        cell_1.font = fonte_tamanho_12_cinza
        cell_1.fill = fundo_cinza

        linha_atual += 1
        ws.merge_cells(start_row=linha_atual, start_column=1, end_row=linha_atual, end_column=3)
        cell_2 = ws.cell(row=linha_atual, column=1, value=f"{item['descricao_detalhada']}")
        cell_2.font = fonte_tamanho_10

        linha_atual += 1
        ws.cell(row=linha_atual, column=1, value=f"UF: {str(item['unidade'])}").font = fonte_tamanho_10
        ws.cell(row=linha_atual, column=2, value=f"Marca: {item['marca_fabricante']}").font = fonte_tamanho_10
        ws.cell(row=linha_atual, column=3, value=f"Modelo: {item['modelo_versao']}").font = fonte_tamanho_10

        linha_atual += 1
 
        # Removendo o sufixo ".0" de quantidade, se presente
        quantidade_formatada = str(int(item['quantidade'])) if item['quantidade'] == int(item['quantidade']) else str(item['quantidade'])

        # Aplicando a formatação corrigida na célula
        ws.cell(row=linha_atual, column=1, value=f"Quantidade: {quantidade_formatada}").font = fonte_tamanho_10

        
        # Convertendo valores para número antes de formatar
        valor_unitario = float(item['valor_homologado_item_unitario'])
        valor_total = float(item['valor_homologado_total_item'])
        
        # Valor Unitário
        print(f"Processando Valor Unitário para o item {item['item_num']}")
        valor_unitario_formatado = format_currency(valor_unitario)
        ws.cell(row=linha_atual, column=2, value=f"Valor Unitário: {valor_unitario_formatado}").font = fonte_tamanho_10
        
        # Valor Total
        print(f"Processando Valor Total para o item {item['item_num']}")
        valor_total_formatado = format_currency(valor_total)
        ws.cell(row=linha_atual, column=3, value=f"Valor Total: {valor_total_formatado}").font = fonte_tamanho_10

        linha_atual += 1

    wb.save(caminho_arquivo_excel)
