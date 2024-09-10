from PyQt6.QtWidgets import *
from PyQt6.QtCore import *
from PyQt6.QtGui import *
from modules.atas.regex_termo_homolog import *
from modules.atas.regex_sicaf import *
from modules.atas.canvas_gerar_atas import criar_pastas_com_subpastas, abrir_pasta, gerar_soma_valor_homologado, inserir_relacao_empresa, inserir_relacao_itens, adicione_texto_formatado
from diretorios import *
import pandas as pd
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate
from datetime import datetime
from modules.planejamento.utilidades_planejamento import DatabaseManager
from docx import Document

class AtasDialog(QDialog):
    NUMERO_ATA_GLOBAL = None  # Defina isso em algum lugar adequado dentro de sua classe
    dataframe_updated = pyqtSignal(object)  # Sinal para emitir o DataFrame atualizado
    pastas_criadas = set()  # Rastreador de pastas criadas

    def __init__(self, parent=None, pe_pattern=None, dataframe=None):
        super().__init__(parent)
        self.db_manager = DatabaseManager(CONTROLE_DADOS)
        self.pe_pattern = pe_pattern
        self.nup_data = None
        self.dataframe = dataframe 
        self.settings = QSettings("YourCompany", "YourApp")  # Ajuste esses valores para seu aplicativo
        self.configurar_ui()

    def closeEvent(self, event):
        # Quando o diálogo é fechado, emite o DataFrame atualizado
        self.dataframe_updated.emit(self.dataframe)
        super().closeEvent(event)

    def update_title_label(self):
        header_layout = QHBoxLayout()

        pixmap_acanto = QPixmap(str(IMAGE_PATH / "acanto"))
        pixmap_acanto = pixmap_acanto.scaled(100, 100, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
        image_label_acanto = QLabel()
        image_label_acanto.setPixmap(pixmap_acanto)
        header_layout.addWidget(image_label_acanto)

        html_text = (
            f"<span style='font-size: 28px'>Painel de Geração de Atas/Contratos</span>"
        )
        self.titleLabel.setText(html_text)
        self.titleLabel.setTextFormat(Qt.TextFormat.RichText)
        self.titleLabel.setStyleSheet(" font-size: 32px; font-weight: bold;")

        header_layout.addSpacerItem(QSpacerItem(40, 20, QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum))
        header_layout.addWidget(self.titleLabel)
        header_layout.addSpacerItem(QSpacerItem(40, 20, QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum))

        image_label_acanto2 = QLabel()
        image_label_acanto2.setPixmap(pixmap_acanto)
        header_layout.addWidget(image_label_acanto2)

        return header_layout

    def apply_widget_style(self, widget):
        widget.setStyleSheet("font-size: 14pt;") 

    def create_button(self, text, icon=None, callback=None, tooltip_text="", button_size=None, icon_size=None):
        btn = QPushButton(text, self)
        if icon:
            btn.setIcon(icon)
            btn.setIconSize(icon_size if icon_size else QSize(40, 40))
        if button_size:
            btn.setFixedSize(button_size)
        btn.setToolTip(tooltip_text)
        btn.setFont(QFont('Arial', 14))
        if callback:
            btn.clicked.connect(callback)
        return btn

    def salvar_alteracoes(self):
        texto_header = self.header_editor.toHtml()  # Captura o texto HTML do QTextEdit
        salvar_configuracoes({'header_text': texto_header})

    def configurar_ui(self):
        self.setWindowTitle("Atas / Contratos")
        self.setFont(QFont('Arial', 14))

        # Cria um QVBoxLayout principal para o QDialog
        layout = QVBoxLayout(self)
        self.setFixedSize(1000, 650) 

        # Cria e configura o titleLabel
        self.titleLabel = QLabel()
        header_layout = self.update_title_label()
        layout.addLayout(header_layout)

        header_label = QLabel("Editar Cabeçalho:")
        header_label.setFont(QFont('Arial', 14))
        layout.addWidget(header_label)
        
        self.header_editor = QTextEdit()
        self.header_editor.setFont(QFont('Arial', 12))
        self.header_editor.setMinimumHeight(180)

        configuracoes = carregar_configuracoes()
        initial_text = configuracoes.get('header_text', '')  # Carrega o texto salvo ou usa string vazia
        self.header_editor.setHtml(initial_text)
        layout.addWidget(self.header_editor)
                
        # Texto inicial com HTML para formatação
        initial_text = ("A União, por intermédio do CENTRO DE INTENDÊNCIA DA MARINHA EM BRASÍLIA (CeIMBra), com sede na "
                        "Esplanada dos Ministérios, Bloco “N”, Prédio Anexo, 2º andar, CEP: 70055-900, na cidade de Brasília – DF, "
                        "inscrito(a) sob o CNPJ nº 00.394.502/0594-67, neste ato representado pelo Capitão de Fragata (IM) "
                        "Thiago Martins Amorim, Ordenador de Despesa, nomeado(a) pela Portaria nº 241 de 25 de abril de 2024, "
                        "do Com7°DN, c/c Ordem de Serviço nº 57/2024 de 25 de abril de 2024 do CeIMBra, considerando o "
                        "julgamento da licitação na modalidade de pregão, na forma eletrônica, para REGISTRO DE PREÇOS nº "
                        "<span style='color: yellow;'>{{num_pregao}}</span>/2024, processo administrativo nº <span style='color: yellow;'>{{nup}}</span>, RESOLVE registrar os preços da(s) "
                        "empresa(s) indicada(s) e qualificada(s) nesta ATA, de acordo com a classificação por ela(s) alcançada(s) "
                        "e na(s) quantidade(s) cotada(s), atendendo as condições previstas no Edital de licitação, sujeitando-se "
                        "as partes às normas constantes na Lei nº 14.133, de 1º de abril de 2021, no Decreto n.º 11.462, de "
                        "31 de março de 2023, e em conformidade com as disposições a seguir:")
        self.header_editor.setHtml(initial_text)
        layout.addWidget(self.header_editor)
        layout.addItem(QSpacerItem(20, 10, QSizePolicy.Policy.Minimum, QSizePolicy.Policy.Fixed))

        # Configurar o combobox das cidades
        cidades_label = QLabel("Selecione a Cidade:")
        cidades_label.setFont(QFont('Arial', 14))
        cidades_label.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Fixed)
        layout.addWidget(cidades_label)

        self.cidades_combobox = QComboBox()
        cidades = ["Brasília-DF", "Rio Grande-RS", "Salvador-BA", "São Pedro da Aldeia-RJ", "Rio de Janeiro-RJ", "Natal-RN", "Manaus-MA"]
        self.cidades_combobox.addItems(cidades)
        self.cidades_combobox.setFont(QFont('Arial', 14))
        self.cidades_combobox.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
        self.cidades_combobox.setMaximumWidth(300)  # Ajusta a largura máxima do combobox

        # Carregar a última cidade selecionada
        last_city = self.settings.value("last_selected_city", "Brasília-DF")
        index = self.cidades_combobox.findText(last_city)
        if index != -1:
            self.cidades_combobox.setCurrentIndex(index)
        
        layout.addWidget(self.cidades_combobox)

        layout.addItem(QSpacerItem(20, 10, QSizePolicy.Policy.Minimum, QSizePolicy.Policy.Fixed))

        # Configuração para seleção da organização gerenciadora
        org_label = QLabel("Selecione a Organização Gerenciadora:")
        org_label.setFont(QFont('Arial', 14))
        layout.addWidget(org_label)

        self.org_combobox = QComboBox()
        organizations = [
            "Centro de Intendência da Marinha em Brasília (CeIMBra)",
            "Centro de Intendência da Marinha em Natal (CeIMNa)",
            "Centro de Intendência da Marinha em Manaus (CeIMMa)",
            "Centro de Intendência da Marinha em Salvador (CeIMSa)",
            "Centro de Intendência da Marinha em Rio Grande (CeIMRG)",
            "Centro de Intendência da Marinha em São Pedro da Aldeia (CeIMSPA)",
            "Hospital Naval de Brasília (HNBra)",
            "Hospital Naval Marcílio Dias (HNMD)",
            "Banco Central do Brasil (BACEN-BSB)",
            "Banco Central do Brasil (BACEN-RJ)",
        ]
        self.org_combobox.addItems(organizations)
        self.org_combobox.setFont(QFont('Arial', 14))
        self.org_combobox.setMaximumWidth(700)
        last_org = self.settings.value("last_selected_org", organizations[0])
        self.org_combobox.setCurrentText(last_org)
        layout.addWidget(self.org_combobox)

        layout.addItem(QSpacerItem(20, 10, QSizePolicy.Policy.Minimum, QSizePolicy.Policy.Fixed))

        self.configurar_rotulos(layout)
        # self.configurar_entrada_e_botao_confirmar(layout)
        # # self.configurar_botoes_acao(layout)
        # self.carregar_e_exibir_ultimo_contrato()
        self.criar_botao_gerar_ata(layout)
        self.setLayout(layout)

    def closeEvent(self, event):
        # Salvar as últimas seleções
        self.settings.setValue("last_selected_city", self.cidades_combobox.currentText())
        self.settings.setValue("last_selected_org", self.org_combobox.currentText())
        self.dataframe_updated.emit(self.dataframe)
        super().closeEvent(event)

    def configurar_rotulos(self, layout):
        rotulo = QLabel("Digite o próximo Número de Controle de Atas/Contratos:")
        rotulo.setFont(QFont('Arial', 14))

        # Criar um QHBoxLayout para o rótulo e o campo de entrada
        linha_rotulo_entrada = QHBoxLayout()
        linha_rotulo_entrada.addWidget(rotulo)

        self.entradaAta = QLineEdit(self)
        self.entradaAta.setFont(QFont('Arial', 14))
        self.entradaAta.setPlaceholderText("Digite um número até 4 dígitos")
        self.entradaAta.setMaxLength(10)

        # Ajustar largura para acomodar mais caracteres, aqui adicionamos mais 50 pixels para dar mais espaço
        self.entradaAta.setFixedWidth(self.entradaAta.fontMetrics().horizontalAdvance('0000') + 250)

        linha_rotulo_entrada.addWidget(self.entradaAta)

        # Usando a função create_button para criar o botão de confirmar
        self.botaoConfirmar = self.create_button("Confirmar", None, self.confirmar_numero_ata_e_nup_do_processo, "Clique para confirmar o número")
        linha_rotulo_entrada.addWidget(self.botaoConfirmar)

        # Adiciona o layout horizontal ao layout vertical principal
        layout.addLayout(linha_rotulo_entrada)
        
    def criar_botao_gerar_ata(self, layout):
        # Caminhos para os ícones
        icon_confirm = QIcon(str(ICONS_DIR / "production_red.png"))  # Caminho para o ícone de confirmação
        
        # Criação do botão
        button_confirm = self.create_button("  Gerar Atas de Registro de Preços", icon_confirm, self.gerar_ata_de_registro_de_precos, "Após inserir as informações, clique para gerar as Atas", QSize(380, 50), QSize(50, 50))
        
        # Criação de um layout horizontal para centralizar o botão
        hbox = QHBoxLayout()
        hbox.addSpacerItem(QSpacerItem(40, 20, QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum))
        hbox.addWidget(button_confirm)
        hbox.addSpacerItem(QSpacerItem(40, 20, QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum))
        
        # Adicionando o layout horizontal ao layout principal
        layout.addLayout(hbox)
        
        # Aplicar estilo ao botão
        self.apply_widget_style(button_confirm)

    @staticmethod
    def convert_pe_format(pe_string):
        if pe_string is None:
            print("Erro: pe_string é None")
            return None

        pe_formatted = pe_string.replace('PE-', 'PE ').replace('-', '/')
        print(f"Converted PE format: {pe_formatted}")  # Depuração
        return pe_formatted

    def obter_nup(self, pe_formatted):
        try:
            with self.db_manager as conn:
                query = f"SELECT nup FROM controle_processos WHERE id_processo LIKE '%{pe_formatted}%'"
                df = pd.read_sql(query, conn)
                if not df.empty:
                    self.nup_data = {
                        'nup': df.iloc[0]['nup']
                    }
                    return self.nup_data
                else:
                    return None
        except Exception as e:
            print(f"Erro ao acessar o banco de dados: {e}")
            return None
                                       
    def atualizar_rotulo_ultimo_contrato(self, ultimo_numero_contrato):
        self.rotulo_ultimo_contrato.setText(f"O último número de ata/contrato gerado foi: {ultimo_numero_contrato}")

    def salvar_ultimo_contrato(self, ultimo_num_contrato):
        with open(ULTIMO_CONTRATO_DIR, "w") as f:
            f.write(str(ultimo_num_contrato))  # Convertendo para string

    def carregar_ultimo_contrato(self):
        try:
            with open(ULTIMO_CONTRATO_DIR, "r") as f:
                return int(f.read().split('/')[-1])
        except (FileNotFoundError, ValueError):
            return None

    def confirmar_numero_ata_e_nup_do_processo(self):
        numero_ata = self.entradaAta.text()
        print(f"Valor inserido pelo usuário: '{numero_ata}'")  # Use aspas para capturar espaços em branco ou strings vazias

        if numero_ata.isdigit() and len(numero_ata) <= 4:
            AtasDialog.NUMERO_ATA_GLOBAL = int(numero_ata)
            
            # Verificação para self.pe_pattern
            if self.pe_pattern is None:
                QMessageBox.warning(self, "Erro", "Padrão de pregão não encontrado. Por favor, carregue um database antes de continuar.")
                print("Padrão de pregão não encontrado. Necessário carregar um database.")
                return
            print(f"Obtido PE: {self.pe_pattern}")  # Depuração
            pe_formatted = self.convert_pe_format(self.pe_pattern)
            print(f"Obtido PE formatado: {pe_formatted}")  # Depuração
            nup = self.obter_nup(pe_formatted)
            print(f"Obtido NUP: {nup}")  # Depuração

            if nup:
                self.nup_data = nup
                QMessageBox.information(self, "Número Confirmado", f"Número da ata definido para: {numero_ata}")
                print(f"Número da ATA confirmado e definido como {numero_ata}.")
            else:
                QMessageBox.warning(self, "Erro ao Obter NUP", "Não foi encontrado um padrão de pregão para prosseguir. Por favor, carregue um database antes de continuar.")
                print("Não foi encontrado um padrão de pregão para prosseguir. Necessário carregar um database.")
        else:
            QMessageBox.warning(self, "Número Inválido", "Por favor, digite um número válido de até 4 dígitos.")
            print("Tentativa de inserção de um número inválido.")


    def gerar_ata_de_registro_de_precos(self):
        if not self.nup_data:  # Verifica se nup_data está vazia ou é None
            self.nup_data = "(INSIRA O NUP)"  # Atribui um valor padrão caso não exista nup_data
        
        # Chama salvar_alteracoes para garantir que todas as mudanças sejam salvas antes de processar
        self.salvar_alteracoes()
        
        self.processar_ata_de_registro_de_precos(self.nup_data, self.dataframe)

    def processar_ata_de_registro_de_precos(self, nup_data, dataframe):
        if AtasDialog.NUMERO_ATA_GLOBAL is None:
            QMessageBox.information(self, "Inserir Número da ATA", "Por favor, insira o número da ATA para continuar.")
            return

        criar_pastas_com_subpastas(dataframe)  # Chamando função externa de criação de pastas
        ultimo_num_ata = self.processar_ata(AtasDialog.NUMERO_ATA_GLOBAL, nup_data, dataframe)

        self.salvar_ultimo_contrato(ultimo_num_ata)
        # self.atualizar_rotulo_ultimo_contrato(ultimo_num_ata)

    def processar_ata(self, NUMERO_ATA, nup_data, dataframe):
        if isinstance(nup_data, dict) and 'nup' in nup_data:
            nup = nup_data['nup']
        else:
            nup = "(INSIRA O NUP)" 
        relatorio_path = get_relatorio_path()
        combinacoes = dataframe[['uasg', 'num_pregao', 'ano_pregao', 'empresa']].drop_duplicates().values
        NUMERO_ATA_atualizado = NUMERO_ATA

        # Verifica se a coluna 'numero_ata' existe, senão, cria ela
        if 'numero_ata' not in dataframe.columns:
            dataframe['numero_ata'] = None

        for uasg, num_pregao, ano_pregao, empresa in combinacoes:
            if not pd.isna(num_pregao) and not pd.isna(ano_pregao) and not pd.isna(empresa):
                path_dir_principal, path_subpasta = self.preparar_diretorios(relatorio_path, num_pregao, ano_pregao, empresa)
                registros_empresa = dataframe[dataframe['empresa'] == empresa]
                NUMERO_ATA_atualizado, num_contrato = self.processar_empresa(registros_empresa, empresa, path_subpasta, nup, NUMERO_ATA_atualizado)

                # Atualizar o DataFrame com o número de contrato atualizado para a empresa processada
                dataframe.loc[dataframe['empresa'] == empresa, 'numero_ata'] = num_contrato

        abrir_pasta(str(path_dir_principal))
        print(dataframe[['numero_ata', 'item_num']])  # Mostra os valores das colunas 'numero_ata' e 'item_num'
        return NUMERO_ATA_atualizado

    def preparar_diretorios(self, relatorio_path, num_pregao, ano_pregao, empresa):        
        nome_dir_principal = f"PE {int(num_pregao)}-{int(ano_pregao)}"
        path_dir_principal = relatorio_path / nome_dir_principal
        nome_empresa_limpa = self.limpar_nome_empresa(empresa)
        path_subpasta = path_dir_principal / nome_empresa_limpa

        # Verifica se a pasta já foi criada anteriormente
        chave_pasta = (nome_dir_principal, nome_empresa_limpa)
        if chave_pasta not in AtasDialog.pastas_criadas:
            if not path_subpasta.exists():
                path_subpasta.mkdir(parents=True, exist_ok=True)
                print(f"Criado subdiretório: {path_subpasta}")
            AtasDialog.pastas_criadas.add(chave_pasta)
        else:
            print(f"Subdiretório já existente, não será recriado: {path_subpasta}")

        return path_dir_principal, path_subpasta

    def limpar_nome_empresa(self, nome_empresa):
        # Substituir '/' e ':' por sublinhado
        caracteres_para_sublinhado = ['/', ':']
        for char in caracteres_para_sublinhado:
            nome_empresa = nome_empresa.replace(char, '_')
        
        # Substituir '.' por nada (remover)
        nome_empresa = nome_empresa.replace('.', '')

        # Substituir outros caracteres inválidos por sublinhados
        caracteres_invalidos = ['<', '>', '_', '"', '\\', '|', '?', '*']
        for char in caracteres_invalidos:
            nome_empresa = nome_empresa.replace(char, '_')

        # Remover espaços extras e sublinhados no final do nome da empresa
        nome_empresa = nome_empresa.rstrip(' _')

        # Substituir múltiplos espaços ou sublinhados consecutivos por um único sublinhado
        nome_empresa = '_'.join(filter(None, nome_empresa.split(' ')))

        return nome_empresa
    
    def processar_empresa(self, registros_empresa, empresa, path_subpasta, nup, NUMERO_ATA_atualizado):
        if not registros_empresa.empty:
            registro = registros_empresa.iloc[0].to_dict()
            itens_relacionados = registros_empresa.to_dict('records')
            context, num_contrato = self.criar_contexto(registro, empresa, NUMERO_ATA_atualizado, nup, itens_relacionados)
            
            # Passando num_contrato para salvar_documento
            self.salvar_documento(path_subpasta, empresa, context, registro, itens_relacionados, num_contrato)

            # Garanta que NUMERO_ATA_atualizado é um inteiro antes de incrementar
            NUMERO_ATA_atualizado = int(NUMERO_ATA_atualizado) + 1  # Atualiza o número da ATA após processar com sucesso
        else:
            print(f"Nenhum registro encontrado para a empresa: {empresa}")
            num_contrato = None
        return NUMERO_ATA_atualizado, num_contrato


    def criar_contexto(self, registro, empresa, NUMERO_ATA_atualizado, nup, itens_relacionados):
        ano_atual = datetime.now().year
        num_contrato = f"{registro['uasg']}/{ano_atual}-{NUMERO_ATA_atualizado:03}/00"
        texto_substituto = f"Pregão Eletrônico nº {registro['num_pregao']}/{registro['ano_pregao']}\n{num_contrato}"
        soma_valor_homologado = gerar_soma_valor_homologado(itens_relacionados)

        dados_da_ug_contratante_cabecalho = self.header_editor.toPlainText()
        cidade_selecionada = self.cidades_combobox.currentText()
        organizacao_selecionada = self.org_combobox.currentText()

        return ({
            "num_pregao": registro['num_pregao'],
            "ano_pregao": registro['ano_pregao'],
            "empresa": empresa,
            "uasg": registro['uasg'],
            "numero_ata": NUMERO_ATA_atualizado,
            "soma_valor_homologado": soma_valor_homologado,
            "cabecalho": texto_substituto,
            "dados_ug_contratante": dados_da_ug_contratante_cabecalho,
            "contrato": num_contrato,
            "endereco": registro["endereco"],
            "cnpj": registro["cnpj"],
            "objeto": registro["objeto"],
            "ordenador_despesa": registro["ordenador_despesa"],
            "responsavel_legal": registro["responsavel_legal"],
            "nup": nup,
            "email": registro["email"],
            "cidade": cidade_selecionada,
            "organizacao": organizacao_selecionada
        }, num_contrato)

    def salvar_email(self, path_subpasta, context):
        nome_arquivo_txt = "E-mail.txt"
        path_arquivo_txt = path_subpasta / nome_arquivo_txt
        with open(path_arquivo_txt, "w") as arquivo_txt:
            texto_email = (f"{context['email']}\n\n"
                        f"Sr. Representante.\n\n"
                        f"Encaminho em anexo a Vossa Senhoria a ATA {context['contrato']} "
                        f"decorrente do Pregão Eletrônico (SRP) nº {context['num_pregao']}/{context['ano_pregao']}, do Centro "
                        f"de Intendência da Marinha em Brasília (CeIMBra).\n\n"
                        f"Os documentos deverão ser conferidos, assinados e devolvidos a este Comando.\n\n"
                        f"A empresa receberá uma via, devidamente assinada, após a publicação.\n\n"
                        f"Respeitosamente,\n")
            arquivo_txt.write(texto_email)

    def salvar_documento(self, path_subpasta, empresa, context, registro, itens_relacionados, num_contrato):
        max_len = 40  # Definindo o limite máximo para o nome da empresa
        contrato_limpo = self.limpar_nome_empresa(num_contrato)[:max_len].rstrip()

        # Preparar o template do documento
        tpl = DocxTemplate(TEMPLATE_PATH)
        tpl.render(context)

        # Montar o nome do arquivo, garantindo que não ultrapasse os limites comuns de sistemas de arquivos
        nome_documento = f"{contrato_limpo}.docx"
        path_documento = path_subpasta / nome_documento

        # Salvar o documento
        tpl.save(path_documento)

        # Alterar o documento após a criação inicial para incluir informações detalhadas
        self.alterar_documento_criado(path_documento, registro, registro["cnpj"], itens_relacionados)

        # Salvando o arquivo de email associado
        self.salvar_email(path_subpasta, context)

    def alterar_documento_criado(self, caminho_documento, registro, cnpj, itens):
        # Carregar o documento
        doc = Document(caminho_documento)

        # Iterar por cada parágrafo
        for paragraph in doc.paragraphs:
            if '{relacao_empresa}' in paragraph.text:
                paragraph.clear()
                inserir_relacao_empresa(paragraph, registro, cnpj)

            if '{relacao_item}' in paragraph.text:
                paragraph.clear()
                inserir_relacao_itens(paragraph, itens)

        # Extrair o diretório do caminho do documento
        diretorio_documento = os.path.dirname(caminho_documento)

        # Gerar o arquivo Excel no mesmo diretório do documento
        caminho_arquivo_excel = os.path.join(diretorio_documento, 'relacao_itens.xlsx')
        gerar_excel_relacao_itens(itens, caminho_arquivo_excel)

        # Inserir a tabela do Excel no local do marcador <<tabela_itens>>
        inserir_tabela_do_excel_no_documento(doc, caminho_arquivo_excel)

        # Salvar o documento atualizado
        doc.save(caminho_documento)

    def inserir_relacao_empresa(self, paragrafo, registro, cnpj):
        dados = {
            "Razão Social": registro["empresa"],
            "CNPJ": registro["cnpj"],
            "Endereço": registro["endereco"],
            "Município-UF": registro["municipio"],
            "CEP": registro["cep"],
            "Telefone": registro["telefone"],
            "E-mail": registro["email"]
        }

        total_itens = len(dados)
        contador = 1
        
        for chave, valor in dados.items():
            adicione_texto_formatado(paragrafo, f"{chave}: ", True)

            # Verifica se é a penúltima linha
            if contador == total_itens - 1:
                adicione_texto_formatado(paragrafo, f"{valor}; e\n", False)
            # Verifica se é a última linha
            elif contador == total_itens:
                adicione_texto_formatado(paragrafo, f"{valor}.\n", False)
            else:
                adicione_texto_formatado(paragrafo, f"{valor};\n", False)

            contador += 1
        
        adicione_texto_formatado(paragrafo, "Representada neste ato, por seu representante legal, o(a) Sr(a) ", False)
        adicione_texto_formatado(paragrafo, f'{registro["responsavel_legal"]}.\n', False)

from openpyxl import load_workbook
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def inserir_tabela_do_excel_no_documento(doc, caminho_arquivo_excel):
    # Carregar a planilha do Excel
    wb = load_workbook(caminho_arquivo_excel)
    ws = wb.active

    for paragraph in doc.paragraphs:
        if '<<tabela_itens>>' in paragraph.text:
            # Remover o texto do marcador
            paragraph.text = paragraph.text.replace('<<tabela_itens>>', '')

            # Criar uma nova tabela no documento Word
            table = doc.add_table(rows=0, cols=3)

            for i in range(0, ws.max_row, 4):
                # Primeira linha (mesclada e pintada de cinza claro)
                row_cells = table.add_row().cells
                row_cells[0].merge(row_cells[2])
                run = row_cells[0].paragraphs[0].add_run(str(ws.cell(row=i+1, column=1).value))
                run.font.size = Pt(12)
                run.font.bold = True
                shading_elm = parse_xml(r'<w:shd {} w:fill="D3D3D3"/>'.format(nsdecls('w')))
                row_cells[0]._element.get_or_add_tcPr().append(shading_elm)

                # Segunda linha (mesclada com "Descrição Detalhada:" em negrito e quebras de linha)
                row_cells = table.add_row().cells
                row_cells[0].merge(row_cells[2])
                
                # Adiciona quebra de linha antes de "Descrição Detalhada:"
                run = row_cells[0].paragraphs[0].add_run("\n")
                
                # Adiciona "Descrição Detalhada:" em negrito
                run = row_cells[0].paragraphs[0].add_run("Descrição Detalhada:")
                run.font.size = Pt(10)
                run.font.bold = True

                # Adicionando o texto restante após "Descrição Detalhada:"
                texto_segunda_linha = str(ws.cell(row=i+2, column=1).value)
                run = row_cells[0].paragraphs[0].add_run(f" {texto_segunda_linha}")
                run.font.size = Pt(10)
                
                # Adiciona quebra de linha após o texto
                row_cells[0].paragraphs[0].add_run("\n")

                # Terceira linha (manter formatação padrão)
                row_cells = table.add_row().cells
                for j in range(3):
                    value = ws.cell(row=i+3, column=j+1).value
                    if value is not None:
                        texto = str(value)
                        if j == 0 and texto.startswith("UF:"):
                            # Negrito apenas para "UF:"
                            run = row_cells[j].paragraphs[0].add_run("UF:")
                            run.font.size = Pt(10)
                            run.font.bold = True
                            
                            # Texto que segue "UF:"
                            run = row_cells[j].paragraphs[0].add_run(texto[3:])
                            run.font.size = Pt(10)
                        elif j == 1 and texto.startswith("Marca:"):
                            # Negrito apenas para "Marca:"
                            run = row_cells[j].paragraphs[0].add_run("Marca:")
                            run.font.size = Pt(10)
                            run.font.bold = True
                            
                            # Texto que segue "Marca:"
                            run = row_cells[j].paragraphs[0].add_run(texto[6:])
                            run.font.size = Pt(10)
                        elif j == 2 and texto.startswith("Modelo:"):
                            # Negrito apenas para "Modelo:"
                            run = row_cells[j].paragraphs[0].add_run("Modelo:")
                            run.font.size = Pt(10)
                            run.font.bold = True
                            
                            # Texto que segue "Modelo:"
                            run = row_cells[j].paragraphs[0].add_run(texto[7:])
                            run.font.size = Pt(10)
                        else:
                            row_cells[j].text = texto
                            row_cells[j].paragraphs[0].runs[0].font.size = Pt(10)

                # Quarta linha (manter formatação padrão)
                row_cells = table.add_row().cells
                for j in range(3):
                    value = ws.cell(row=i+4, column=j+1).value
                    if value is not None:
                        texto = str(value)
                        if j == 0 and texto.startswith("Quantidade:"):
                            # Negrito apenas para "Quantidade:"
                            run = row_cells[j].paragraphs[0].add_run("Quantidade:")
                            run.font.size = Pt(10)
                            run.font.bold = True

                            # Texto que segue "Quantidade:"
                            run = row_cells[j].paragraphs[0].add_run(texto[11:])
                            run.font.size = Pt(10)
                            row_cells[j].paragraphs[0].add_run("\n")  # Adiciona quebra de linha
                        elif j == 1 and texto.startswith("Valor Unitário:"):
                            # Negrito apenas para "Valor Unitário:"
                            run = row_cells[j].paragraphs[0].add_run("Valor Unitário:")
                            run.font.size = Pt(10)
                            run.font.bold = True

                            # Texto que segue "Valor Unitário:"
                            run = row_cells[j].paragraphs[0].add_run(texto[15:])
                            run.font.size = Pt(10)
                            row_cells[j].paragraphs[0].add_run("\n")  # Adiciona quebra de linha
                        elif j == 2 and texto.startswith("Valor Total:"):
                            # Negrito apenas para "Valor Total:"
                            run = row_cells[j].paragraphs[0].add_run("Valor Total:")
                            run.font.size = Pt(10)
                            run.font.bold = True

                            # Texto que segue "Valor Total:"
                            run = row_cells[j].paragraphs[0].add_run(texto[12:])
                            run.font.size = Pt(10)
                            row_cells[j].paragraphs[0].add_run("\n")  # Adiciona quebra de linha
                        else:
                            row_cells[j].text = texto
                            row_cells[j].paragraphs[0].runs[0].font.size = Pt(10)
                            row_cells[j].paragraphs[0].add_run("\n")  # Adiciona quebra de linha
            # Mover a tabela para logo após o parágrafo atual
            move_table_after_paragraph(paragraph, table)
            break

def move_table_after_paragraph(paragraph, table):
    # Move a tabela para ficar logo após o parágrafo atual
    tbl, p = table._tbl, paragraph._element
    p.addnext(tbl)

def salvar_configuracoes(dados):
    with open('configuracoes.json', 'w') as arquivo:
        json.dump(dados, arquivo)

def carregar_configuracoes():
    try:
        with open('configuracoes.json', 'r') as arquivo:
            return json.load(arquivo)
    except (FileNotFoundError, json.JSONDecodeError):
        return {}  # Retorna um dicionário vazio se o arquivo não existir ou estiver corrompido
    
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill

def gerar_excel_relacao_itens(itens, caminho_arquivo_excel='relacao_itens.xlsx'):
    # Ordenar os itens, primeiro por 'grupo' (None será considerado menor) e depois por 'item_num'
    itens_ordenados = sorted(itens, key=lambda x: (x['grupo'] if x['grupo'] is not None else '', x['item_num']))
    
    # Print do ordenamento
    for item in itens_ordenados:
        print(f"Grupo: {item['grupo']} - Item Num: {item['item_num']}")

    wb = Workbook()
    ws = wb.active

    fonte_tamanho_12_cinza = Font(size=12, bold=True)
    fundo_cinza = PatternFill(start_color="D3D3D3", end_color="D3D3D3", fill_type="solid")
    fonte_tamanho_10 = Font(size=10)

    linha_atual = 1

    for item in itens_ordenados:
        ws.merge_cells(start_row=linha_atual, start_column=1, end_row=linha_atual, end_column=3)
        
        # Adicionando lógica para verificar o valor de grupo
        if isinstance(item['grupo'], (int, float)):  # Verifica se 'grupo' é um número
            valor_cell_1 = f"Grupo {item['grupo']} - Item {item['item_num']} - {item['descricao_tr']} ({item['catalogo']})"
        else:
            valor_cell_1 = f"Item {item['item_num']} - {item['descricao_tr']} ({item['catalogo']})"
        
        
        cell_1 = ws.cell(row=linha_atual, column=1, value=valor_cell_1)
        cell_1.font = fonte_tamanho_12_cinza
        cell_1.fill = fundo_cinza

        linha_atual += 1
        ws.merge_cells(start_row=linha_atual, start_column=1, end_row=linha_atual, end_column=3)
        cell_2 = ws.cell(row=linha_atual, column=1, value=f"{item['descricao_detalhada']}")
        cell_2.font = fonte_tamanho_10

        linha_atual += 1
        ws.cell(row=linha_atual, column=1, value=f"UF: {str(item['unidade'])}").font = fonte_tamanho_10
        ws.cell(row=linha_atual, column=2, value=f"Marca: {item['marca_fabricante']}").font = fonte_tamanho_10
        ws.cell(row=linha_atual, column=3, value=f"Modelo: {item['modelo_versao']}").font = fonte_tamanho_10

        linha_atual += 1
 
        # Removendo o sufixo ".0" de quantidade, se presente
        quantidade_formatada = str(int(item['quantidade'])) if item['quantidade'] == int(item['quantidade']) else str(item['quantidade'])

        # Aplicando a formatação corrigida na célula
        ws.cell(row=linha_atual, column=1, value=f"Quantidade: {quantidade_formatada}").font = fonte_tamanho_10

        # Convertendo valores para número antes de formatar
        valor_unitario = float(item['valor_homologado_item_unitario'])
        valor_total = float(item['valor_homologado_total_item'])
        
        # Valor Unitário
        print(f"Processando Valor Unitário para o item {item['item_num']}")
        valor_unitario_formatado = format_currency(valor_unitario)
        ws.cell(row=linha_atual, column=2, value=f"Valor Unitário: {valor_unitario_formatado}").font = fonte_tamanho_10
        
        # Valor Total
        print(f"Processando Valor Total para o item {item['item_num']}")
        valor_total_formatado = format_currency(valor_total)
        ws.cell(row=linha_atual, column=3, value=f"Valor Total: {valor_total_formatado}").font = fonte_tamanho_10

        linha_atual += 1

    wb.save(caminho_arquivo_excel)

def format_currency(value):
    """ Função para formatar valores monetários no formato brasileiro. """
    print(f"Valor original: {value}")  # Print para depuração

    if isinstance(value, str):
        # Tentar converter string para float
        try:
            value = float(value.replace('.', '').replace(',', '.'))
        except ValueError:
            return value  # Se a conversão falhar, retorna o valor original
    elif not isinstance(value, (int, float)):
        return value  # Retorna o valor original se não for um número

    # Formatar o valor como moeda brasileira
    formatted_value = f"R$ {value:,.2f}".replace(',', 'temp').replace('.', ',').replace('temp', '.')
    
    print(f"Valor formatado: {formatted_value}")  # Print para depuração
    return formatted_value