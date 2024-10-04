from docx import Document
from openpyxl import load_workbook
from docx.shared import Pt, Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from pathlib import Path
import pandas as pd
import os
from modules.matriz_de_riscos.mapa_calor import HeatmapGenerator
from matplotlib.colors import LinearSegmentedColormap, to_hex

class TabelaDeRiscos:
    def __init__(self, template_matriz_riscos, template_matriz_parte2, dados):
        self.template_matriz_riscos = template_matriz_riscos
        self.template_matriz_parte2 = template_matriz_parte2
        self.dados = dados
        self.excel_file = 'tabela_de_riscos.xlsx'
        self.output_file = 'Tabela_de_Riscos_Preenchida.docx'
        self.cmap = HeatmapGenerator.create_custom_cmap()

    def criar_tabela_excel(self):
        # Cria um DataFrame
        df_riscos = pd.DataFrame(self.dados)
        
        # Calcula a coluna (P) * (I)
        df_riscos['P*I'] = df_riscos['P'] * df_riscos['I']
        
        # Salva o DataFrame em um arquivo Excel
        df_riscos.to_excel(self.excel_file, index=False)
        print(f"Tabela de riscos salva em {self.excel_file}")

    def carregar_template_word(self):
        # Carrega o template do documento
        self.doc = Document(self.template_matriz_riscos)
        print(f"Template loaded: {self.template_matriz_riscos}")

    def inserir_quebra_pagina(self):
        # Adiciona uma quebra de página após a primeira página
        self.doc.add_page_break()
        print("Quebra de página inserida")

    def configurar_paisagem(self):
        # Adiciona uma nova seção em paisagem após a quebra de página
        new_section = self.doc.add_section(WD_SECTION.NEW_PAGE)
        new_section.orientation = WD_ORIENT.LANDSCAPE
        new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
        
        # Definindo margens estreitas
        new_section.top_margin = Inches(0.2)
        new_section.bottom_margin = Inches(0.2)
        new_section.left_margin = Inches(0.2)
        new_section.right_margin = Inches(0.2)

    def inserir_titulo(self, titulo):
        titulo_paragrafo = self.doc.add_paragraph()
        titulo_run = titulo_paragrafo.add_run(titulo)
        titulo_run.bold = True
        titulo_run.font.size = Pt(14)
        try:
            titulo_run.font.name = 'Carlito'
        except ValueError:
            titulo_run.font.name = 'Calibri'

    def get_fill_color(self, value):
        if value is None:
            return "FFFFFF"  # Retorna branco se o valor for None
        try:
            normalized_value = float(value) / 25  # Normalizar para a escala de 0 a 1
        except ValueError:
            normalized_value = 0  # Se não puder converter, usar 0 como valor padrão
        color = to_hex(self.cmap(normalized_value))
        return color
    
    def inserir_tabela(self, df):
        temp_excel_file = 'temp_tabela_de_riscos.xlsx'
        df.to_excel(temp_excel_file, index=False)
        wb = load_workbook(temp_excel_file)
        ws = wb.active
        print("Tabela carregada do Excel")
        table = self.doc.add_table(rows=ws.max_row, cols=ws.max_column)
        print(f"Tabela criada no Word com {ws.max_row} linhas e {ws.max_column} colunas")
        
        col_widths = [Inches(0.50), Inches(2.50), Inches(2.50), Inches(4.50), Inches(0.3), Inches(0.3), Inches(0.6)]
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            for j, value in enumerate(row):
                cell = table.cell(i, j)
                cell.text = str(value)
                cell_font = cell.paragraphs[0].runs[0].font
                cell_font.size = Pt(10)
                try:
                    cell_font.name = 'Carlito'
                except ValueError:
                    cell_font.name = 'Calibri'
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell_borders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/><w:left w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/><w:right w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/></w:tcBorders>')
                cell._tc.get_or_add_tcPr().append(cell_borders)
                if i == 0:
                    cell_font.bold = True
                    shading_elm = parse_xml(r'<w:shd {} w:fill="BFBFBF"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                elif i % 2 == 1:
                    shading_elm = parse_xml(r'<w:shd {} w:fill="E7E6E6"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)

            if i % 2 == 1:
                for j, cell in enumerate(table.row_cells(i)):
                    shading_elm = parse_xml(r'<w:shd {} w:fill="E7E6E6"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)

        for col_num, width in enumerate(col_widths):
            for row_num in range(ws.max_row):
                cell = table.cell(row_num, col_num)
                cell.width = width
                print(f"Largura da coluna {col_num+1} ajustada para {width}")

        # Colorir a coluna 'P*I'
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            cell = table.cell(i, 6)
            value = row[6]
            fill_color = self.get_fill_color(value)
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), fill_color[1:]))
            cell._tc.get_or_add_tcPr().append(shading_elm)
            print(f"Coluna 'P*I' da linha {i} ajustada com a cor: {fill_color}")

        # Colorir a linha índice
        for i in range(ws.max_row):
            if i == 0:  # Substitua 'index_to_color' pelo índice da linha que deseja colorir
                for j in range(ws.max_column):
                    cell = table.cell(i, j)
                    shading_elm = parse_xml(r'<w:shd {} w:fill="CCCCCC"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                print(f"Linha índice {i} ajustada com a cor amarela.")

        os.remove(temp_excel_file)
        
    def configurar_retrato(self):
        self.doc.add_page_break()
        print("Quebra de página inserida após a tabela")
        new_section = self.doc.add_section(WD_SECTION.NEW_PAGE)
        new_section.orientation = WD_ORIENT.PORTRAIT
        new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
        print("Seção em retrato configurada")

    def inserir_conteudo_template(self):
        template_doc = Document(self.template_matriz_parte2)
        for element in template_doc.element.body:
            self.doc.element.body.append(element)
        print("Conteúdo do template parte 2 inserido")

    def salvar_documento(self):
        self.doc.save(self.output_file)
        print(f"Documento final salvo como {self.output_file}")
        os.startfile(self.output_file)
        print("Documento final aberto")

    def gerar_documento(self):
        self.criar_tabela_excel()
        self.carregar_template_word()
        self.configurar_paisagem()
        fases = {
            "Planejamento da Contratação": [],
            "Seleção do Fornecedor": [],
            "Gestão e Fiscalização do Contrato": []
        }
        for dado in self.dados:
            fase = dado.get("Fase")
            if fase in fases:
                fases[fase].append(dado)
        for titulo, dados_etapa in fases.items():
            if dados_etapa:
                self.inserir_titulo(titulo)
                df_etapa = pd.DataFrame(dados_etapa).drop(columns=["Fase"])
                df_etapa['P*I'] = df_etapa['P'] * df_etapa['I']
                self.inserir_tabela(df_etapa)
                self.doc.add_paragraph()
        heatmap_generator = HeatmapGenerator(pd.DataFrame(self.dados))
        image_path = heatmap_generator.generate_heatmap()
        self.doc.add_picture(image_path, width=Inches(11))
        print(f"Imagem de heatmap {image_path} inserida no documento")
        self.inserir_quebra_pagina()
        self.configurar_retrato()
        self.inserir_conteudo_template()
        self.salvar_documento()

# Utilização da classe
if __name__ == "__main__":
    BASE_DIR = Path(__file__).resolve().parent
    TEMPLATE_MATRIZ_RISCOS = BASE_DIR / "template_matriz_riscos.docx"
    TEMPLATE_MATRIZ_PARTE2 = BASE_DIR / "template_matriz_parte2.docx"
    
    dados = [
        {"Fase": "Planejamento da Contratação","Risco": "R1", "Causa": "Definição de requisitos da contratação insuficientes ou indevidos.",
         "Evento": "Baixa participação/adesão ao registro de preços.", "Consequência": "Risco de imagem para a Central; não atingimento dos objetivos de centralização dos procedimentos de licitação e de padronização da estratégia da contratação, resultando em perdas de economia de escala, visto a baixa participação dos órgãos e entidades da APF.", "P": 1, "I": 5},
        {"Fase": "Planejamento da Contratação", "Risco": "R2", "Causa": "Definição de requisitos da contratação insuficientes ou indevidos.",
         "Evento": "Contratação de solução que não atende à necessidade que originou a contratação.", "Consequência": "Mau uso de recursos públicos; ineficácia da prestação dos serviços e problemas de gerenciamento e fiscalização dos contratos advindos da licitação.", "P": 1, "I": 5},
        {"Fase": "Planejamento da Contratação", "Risco": "R3", "Causa": "Estimativa da quantidade maior ou menor que a necessidade.",
         "Evento": "Exaurimento do quantitativo da ata antecipado, nos casos de subdimensionamento da necessidade ou de finalização da ata com grande saldo, nos casos de superdimensionamento.", "Consequência": "Realização de novo procedimento de registro de preços antes do prazo programado para os casos de subdimensionamento dos quantitativos; Frustração do mercado e preços não condizentes com a expectativa criada nos fornecedores, nos casos de superdimensionamento.", "P": 3, "I": 4},
        {"Fase": "Seleção do Fornecedor", "Risco": "R4", "Causa": "Não parcelar a solução cujo parcelamento é viável.",
         "Evento": "Restrição à competitividade, principalmente das empresas de pequeno porte. Questionamentos dos órgãos de controle sobre o não parcelamento.", "Consequência": "Aumento dos valores contratados; impugnações ao certame; paralisações do certame advindas das diligências de órgãos de controles externos.", "P": 2, "I": 5},
        {"Fase": "Planejamento da Contratação", "Risco": "R5", "Causa": "Imposição de indicativo de economicidade mínima de 10%.",
         "Evento": "Falta de assertividade quanto à potencialidade de economia e viabilidade de cumprimento do indicativo. Dificuldade na análise da economia gerada após implantação da solução.", "Consequência": "Não atendimento das expectativas de economia de recursos públicos.", "P": 5, "I": 5},
        {"Fase": "Planejamento da Contratação", "Risco": "R6", "Causa": "Coleta insuficiente de preços ou falha de método para realizar a estimativa.",
         "Evento": "Estimativas de custos inadequadas.", "Consequência": "Utilização de parâmetro inadequado para análise da viabilidade da contratação; possibilidade de contratação por preços superfaturados ou ocorrência de deserção e dificuldade de justificar as estimativas.", "P": 1, "I": 5},
        {"Fase": "Planejamento da Contratação", "Risco": "R7", "Causa": "Falta de abrangência da análise de viabilidade da contratação.",
         "Evento": "Não consideração de todos os aspectos necessários à análise de viabilidade da contratação.", "Consequência": "Certame fracassado ou contratação de fornecedor que não é capaz de entregar a solução ou solução que não produz os resultados necessários ao atendimento da demanda.", "P": 1, "I": 5},
        {"Fase": "Planejamento da Contratação", "Risco": "R8", "Causa": "Declaração imprecisa do objeto.",
         "Evento": "Compreensão imprecisa da descrição, quantidade ou prazo.", "Consequência": "Contratação que não atenda à necessidade da organização.", "P": 1, "I": 5},
        {"Fase": "Planejamento da Contratação", "Risco": "R9", "Causa": "Declaração imprecisa do objeto.",
         "Evento": "Inconformidade legal do edital.", "Consequência": "Impugnações ao edital; declaração de nulidade dos procedimentos; responsabilização de agente(s) de contratação e/ou gestores.", "P": 1, "I": 5},
        {"Fase": "Planejamento da Contratação", "Risco": "R10", "Causa": "Definição de mecanismos que propiciem a ingerência da contratante na administração da contratada.",
         "Evento": "Caracterização de execução indireta ilegal.", "Consequência": "Prática de ilícito trabalhista ante os entendimentos contidos na Súmula nº 331/TST.", "P": 1, "I": 3},
        {"Fase": "Planejamento da Contratação", "Risco": "R11", "Causa": "Subjetividade na definição dos resultados que serão mensurados para fins de remuneração da contratada.",
         "Evento": "Pagamentos sem que tenham sido realmente entregues resultados que atendem às necessidades da organização e/ou Pagamentos aquém do resultado atingir pelo fornecedor.", "Consequência": "Desperdício de recursos públicos e não atendimento das necessidades da organização ou prejuízo financeiro à contratada.", "P": 1, "I": 4},
        {"Fase": "Planejamento da Contratação", "Risco": "R12", "Causa": "Empresas sem qualificação econômico-financeira e técnica-operacional para a execução do objeto participando da licitação.",
         "Evento": "Contratação de empresa incapaz de executar o serviço, as obrigações financeiras, fiscais, trabalhistas e previdenciárias relativas ao contrato.", "Consequência": "Rescisão contratual; necessidade de realização de contratação emergencial.", "P": 1, "I": 5},
        {"Fase": "Planejamento da Contratação", "Risco": "R13", "Causa": "Licitante vencedora apresenta proposta com preços de alguns itens abaixo do mercado (subpreço) e de outros itens acima do mercado (sobrepreço), mas de forma que o valor global de sua proposta seja o menor.",
         "Evento": "Contratação de proposta que não vantajosa (jogo de planilhas).", "Consequência": "Dano ao erário em caso de utilização de quantidade maior dos itens com sobrepreço.", "P": 3, "I": 5},
        {"Fase": "Planejamento da Contratação", "Risco": "R14", "Causa": "Utilização como critério de julgamento do menor preço global por grupo de itens (lote).",
         "Evento": "Ata em que o preço registrado global é o mais vantajoso, mas o preço registrado unitário de um ou mais itens pode não ser o menor ou compatível com os preços de mercado.", "Consequência": "Contratação por preços unitários acima do mercado, causando dano ao erário.", "P": 3, "I": 5},
        {"Fase": "Gestão e Fiscalização do Contrato", "Risco": "R15", "Causa": "Responsável pela gestão e fiscalização do contrato não detém as competências multidisciplinares e/ou condições necessárias à execução da atividade.",
         "Evento": "Gestão e/ou fiscalização inadequada.", "Consequência": "Comprometimento do resultado do serviço prestado.", "P": 3, "I": 3},
        {"Fase": "Planejamento da Contratação", "Risco": "R16", "Causa": "Alterações das condições econômico-financeiras do fornecedor.",
         "Evento": "Descumprimento das condições de habilitação e exigidas na licitação.", "Consequência": "Retorno de riscos que foram mitigados por meio dos critérios de habilitação e qualificação da licitação; descontinuidade contratual; pagamento de fornecedor em débito com a fazenda.", "P": 3, "I": 5},
        {"Fase": "Planejamento da Contratação", "Risco": "R17", "Causa": "Falta de sistematização sobre o que deve ser verificado na fiscalização contratual.",
         "Evento": "Aceites provisórios e definitivos em objetos parcialmente executados ou não executados.", "Consequência": "Pagamento indevido e insatisfação dos usuários.", "P": 3, "I": 3},
        {"Fase": "Planejamento da Contratação", "Risco": "R18", "Causa": "Elementos básicos do contrato não estão claros para as partes.",
         "Evento": "Diferenças de entendimentos e de expectativas entre as partes.", "Consequência": "Falhas na execução do contrato.", "P": 3, "I": 3},
        {"Fase": "Planejamento da Contratação", "Risco": "R19", "Causa": "Inadimplência da contratada.",
         "Evento": "Descumprimento das obrigações trabalhistas, previdenciárias e para com o FGTS pela contratada.", "Consequência": "Responsabilização subsidiária da APF em ações judiciais promovidas pelos empregados alocados na execução do contrato; rescisão contratual; necessidade de contratação emergencial.", "P": 3, "I": 5},
        {"Fase": "Planejamento da Contratação", "Risco": "R20", "Causa": "Declaração imprecisa do objeto.",
         "Evento": "Decorrente inadequação dos parâmetros de fiscalização e de gestão contratual definidos no edital e anexos.", "Consequência": "Dificuldade acentuada para a realização da fiscalização e da gestão contratual junto à contratada, mediante os parâmetros exigíveis.", "P": 1, "I": 5}
    ]

    tabela_de_riscos = TabelaDeRiscos(TEMPLATE_MATRIZ_RISCOS, TEMPLATE_MATRIZ_PARTE2, dados)
    tabela_de_riscos.gerar_documento()
