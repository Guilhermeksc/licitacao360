from modules.gerar_atas.regex_termo_homolog import *
from modules.gerar_atas.regex_sicaf import *
import locale
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt
from PyQt6.QtWidgets import QMessageBox
import pandas as pd
from pathlib import Path
import locale
from num2words import num2words
import os
import subprocess
import sys

# Define a localização para o formato de moeda brasileiro
try:
    # Tenta a configuração comum em sistemas baseados em Unix
    locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')
except locale.Error:
    # Tenta a configuração comum em sistemas Windows
    locale.setlocale(locale.LC_ALL, 'Portuguese_Brazil.1252')

# Esta variável global e gerador precisam ser inicializados em algum lugar no início do seu aplicativo PyQt5
NUMERO_ATA_GLOBAL = None  # Deve ser definido em algum ponto antes de iniciar_processo
GERADOR_NUMERO_ATA = None  # Deve ser definido em algum ponto antes de iniciar_processo

def adicione_texto_formatado(paragraph, text, bold=False):
    run = paragraph.add_run(text)
    run.bold = bold
    font = run.font
    font.name = 'Calibri'
    font.size = Pt(12)

def seu_gerador_inicial(valor_inicial: int):
    numero = valor_inicial
    while True:
        valor_recebido = (yield numero)
        if valor_recebido is not None:
            numero = valor_recebido
        else:
            numero += 1

def confirmar_numero_ata(numero_ata, parent=None):
    global NUMERO_ATA_GLOBAL, GERADOR_NUMERO_ATA
    
    NUMERO_ATA_GLOBAL = int(numero_ata)
    mensagem = f'A próxima ata de registro de preços será "{numero_ata}-00"'
    QMessageBox.information(parent, "Confirmação", mensagem)
    
    GERADOR_NUMERO_ATA = seu_gerador_inicial(NUMERO_ATA_GLOBAL)
    next(GERADOR_NUMERO_ATA)

def limpar_nome_empresa(nome_empresa):
    # Substituir '/' e ':' por sublinhado
    nome_empresa = nome_empresa.replace('/', '_').replace(':', '_')
    
    # Substituir '.' por nada (remover)
    nome_empresa = nome_empresa.replace('.', '')

    # Substituir outros caracteres inválidos por sublinhados
    caracteres_invalidos = ['<', '>', '"', '\\', '|', '?', '*']
    for char in caracteres_invalidos:
        nome_empresa = nome_empresa.replace(char, '_')

    # Remover espaços extras e sublinhados no final do nome da empresa
    nome_empresa = nome_empresa.rstrip(' _')

    # Substituir múltiplos espaços ou sublinhados consecutivos por um único sublinhado
    nome_empresa = '_'.join(filter(None, nome_empresa.split(' ')))

    # Remover duplicatas de sublinhados causados por espaços ou caracteres inválidos
    while '__' in nome_empresa:
        nome_empresa = nome_empresa.replace('__', '_')
        
    return nome_empresa.upper()

def criar_diretorio(base_path: Path, num_pregao: int, ano_pregao: int, nome_empresa: str) -> Path:
    nome_dir_principal = f"PE {num_pregao}-{ano_pregao}"
    path_dir_principal = base_path / nome_dir_principal

    if not path_dir_principal.exists():
        path_dir_principal.mkdir(parents=True)
        print(f"Criado diretório principal: {path_dir_principal}")

    nome_empresa_limpa = limpar_nome_empresa(nome_empresa)
    path_subpasta = path_dir_principal / nome_empresa_limpa

    if not path_subpasta.exists():
        path_subpasta.mkdir(parents=True)
        print(f"Criado subdiretório: {path_subpasta}")
    else:
        print(f"O subdiretório já existe e não será recriado: {path_subpasta}")

    return path_subpasta

pastas_criadas = set()

def criar_pastas_com_subpastas(dataframe) -> None:
    if dataframe is None:
        QMessageBox.warning(None, "Erro", "Padrão de pregão não encontrado. Por favor, carregue um database antes de continuar.")
        return
    
    relatorio_path = get_relatorio_path()
    combinacoes = dataframe[['num_pregao', 'ano_pregao', 'empresa']].drop_duplicates().values
    
    pastas_criadas = set()
    
    for num_pregao, ano_pregao, empresa in combinacoes:
        if pd.isna(num_pregao) or pd.isna(ano_pregao) or pd.isna(empresa):
            continue

        chave_pasta = (int(num_pregao), int(ano_pregao), empresa)
        if chave_pasta not in pastas_criadas:
            criar_diretorio(relatorio_path, int(num_pregao), int(ano_pregao), empresa)
            print(f"Criado 1 diretório para {empresa}")
            pastas_criadas.add(chave_pasta)


def abrir_pasta(pasta):
    if sys.platform == "win32":
        os.startfile(pasta)
    elif sys.platform == "darwin":
        subprocess.Popen(["open", pasta])
    else:
        subprocess.Popen(["xdg-open", pasta])

def gerar_soma_valor_homologado(itens):
    valor_total = sum(float(item["valor_homologado_total_item"] or 0) for item in itens)
    valor_extenso = valor_por_extenso(valor_total)
    return f'R$ {formatar_brl(valor_total)} ({valor_extenso})'

def inserir_relacao_empresa(paragrafo, registro, cnpj):
    dados = {
        "Razão Social": registro["empresa"],
        "CNPJ": registro["cnpj"],
        "Endereço": registro["endereco"],
        "Município-UF": registro["municipio"],
        "CEP": registro["cep"],
        "Telefone": registro["telefone"],
        "E-mail": registro["email"]
    }

    total_itens = len(dados)
    contador = 1
    
    for chave, valor in dados.items():
        adicione_texto_formatado(paragrafo, f"{chave}: ", True)
        if contador == total_itens - 1:
            adicione_texto_formatado(paragrafo, f"{valor}; e\n", False)
        elif contador == total_itens:
            adicione_texto_formatado(paragrafo, f"{valor}.\n", False)
        else:
            adicione_texto_formatado(paragrafo, f"{valor};\n", False)
        contador += 1
    
    adicione_texto_formatado(paragrafo, "Representada neste ato, por seu representante legal, o(a) Sr(a) ", False)
    adicione_texto_formatado(paragrafo, f'{registro["responsavel_legal"]}.\n', False)

def validar_e_corrigir_item(item):
    if item["descricao_detalhada"] is None or not item["descricao_detalhada"].strip():
        raise ValueError(f"O campo 'descricao_detalhada' está ausente ou inválido no item: {item['item_num']}")

    return item

    #Padrão Serviços
    # return [
    #     (f'Item {item_num_int}', False),
    #     (f'Descrição: {item["descricao_detalhada_tr"]}', False),
    #     (f'Unidade de Fornecimento: {item["unidade"]}', False),
    #     # (f'Marca/Fabricante: {item["marca_fabricante"]}   |   Modelo/Versão: {item["modelo_versao"]}', False),
    #     (f'Quantidade: {quantidade_formatada}   |   Valor Unitário: R$ {formatar_brl(item["valor_homologado_item_unitario"])}   |   Valor Total do Item: R$ {formatar_brl(item["valor_homologado_total_item"])}', False),
    #     (f'{"-" * 130}', False)
    # ]
    # Padrão Material

def gerar_campos_item(item):
    if item["descricao_detalhada"] is None or not item["descricao_detalhada"].strip():
        raise ValueError(f"O campo 'descricao_detalhada' está ausente ou inválido no item: {item['item_num']}")

    try:
        item_num_int = int(item["item_num"])
        catalogo_int = item["catalogo"]

        quantidade_formatada = f"{float(item['quantidade']):.2f}".rstrip('0').rstrip('.')
        descricao_detalhada_ajustada = item["descricao_detalhada"].replace("\n", " ")

        return [
            (f'Item {item_num_int} - {item["descricao_tr"]} | Catálogo: {catalogo_int}', True),
            (f'Descrição: {descricao_detalhada_ajustada}', False),
            (f'Unidade de Fornecimento: {item["unidade"]}', False),
            (f'Marca/Fabricante: {item["marca_fabricante"]}   |   Modelo/Versão: {item["modelo_versao"]}', False),
            (f'Quantidade: {quantidade_formatada}   |   Valor Unitário: R$ {formatar_brl(item["valor_homologado_item_unitario"])}   |   Valor Total do Item: R$ {formatar_brl(item["valor_homologado_total_item"])}', False),
            (f'{"-" * 130}', False)
        ]

    except ValueError as ve:
        QMessageBox.critical(None, "Erro ao gerar campos do item", str(ve))
        return None  

def inserir_relacao_itens(paragrafo, itens):
    # Limpar parágrafo
    paragrafo.clear()

    # # Iterar sobre os itens e gerar campos
    # for item in itens:
    #     campos = gerar_campos_item(item)
    #     if campos:
    #         for texto, negrito in campos:
    #             adicione_texto_formatado(paragrafo, texto + '\n', negrito)

    # Calcular e adicionar o valor total homologado
    valor_total = sum(float(item["valor_homologado_total_item"] or 0) for item in itens)
    valor_extenso = valor_por_extenso(valor_total)
    texto_soma_valor_homologado = f'{formatar_brl(valor_total)} ({valor_extenso})'
    adicione_texto_formatado(paragrafo, 'Valor total homologado para a empresa:\n', False)
    adicione_texto_formatado(paragrafo, texto_soma_valor_homologado + '\n', True)

    # Chamar a função para gerar o Excel
    gerar_excel_relacao_itens(itens)
    
    return texto_soma_valor_homologado

def gerar_excel_relacao_itens(itens, caminho_arquivo_excel='relacao_itens.xlsx'):
    dados_para_tabela = []

    # Iterar sobre os itens e preparar os dados para o Excel
    for item in itens:
        dados_para_tabela.append({
            'Item Número': item['item_num'],
            'Descrição': item['descricao_tr'],
            'Quantidade': item['quantidade'],
            'Valor Unitário': item['valor_homologado_item_unitario'],
            'Valor Total': item['valor_homologado_total_item']
        })

    # Criar DataFrame e exportar para Excel
    df = pd.DataFrame(dados_para_tabela)
    df.to_excel(caminho_arquivo_excel, index=False)

def formatar_brl(valor):
    try:
        if valor is None:
            return "Não disponível"  # Retorna uma string informativa caso o valor seja None
        # Formata o número no formato monetário brasileiro sem utilizar a biblioteca locale
        return f"R$ {float(valor):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except (ValueError, TypeError):
        return "Valor inválido"  # Retorna isso se não puder converter para float
       
def valor_por_extenso(valor):
    extenso = num2words(valor, lang='pt_BR', to='currency')
    return extenso.capitalize()

def alterar_contrato_criado(caminho_documento, registro, cnpj, itens):
    # Carregando o documento real
    doc = Document(caminho_documento)
    
    # Iterando por cada parágrafo do documento
    for paragraph in doc.paragraphs:
        if '{relacao_empresa}' in paragraph.text:
            # Substituindo o marcador pelo conteúdo gerado pela função inserir_relacao_empresa
            paragraph.clear()  # Limpar o parágrafo atual
            inserir_relacao_empresa(paragraph, registro, cnpj)
        
        # Verificando o marcador {relacao_item}
        if '{relacao_item}' in paragraph.text:
            # Substituindo o marcador pelo conteúdo gerado pela função inserir_relacao_itens
            paragraph.clear()  # Limpar o parágrafo atual
            inserir_relacao_itens(paragraph, itens)
    
    # Salvando as alterações no documento
    doc.save(caminho_documento)
