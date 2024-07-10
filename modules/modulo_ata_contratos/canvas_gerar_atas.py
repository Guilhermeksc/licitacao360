from modules.modulo_ata_contratos.regex_termo_homolog import *
from modules.modulo_ata_contratos.regex_sicaf import *
import locale
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt
from PyQt6.QtWidgets import QMessageBox
import pandas as pd
from pathlib import Path
import locale
from num2words import num2words

# Define a localização para o formato de moeda brasileiro
try:
    # Tenta a configuração comum em sistemas baseados em Unix
    locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')
except locale.Error:
    # Tenta a configuração comum em sistemas Windows
    locale.setlocale(locale.LC_ALL, 'Portuguese_Brazil.1252')

# Esta variável global e gerador precisam ser inicializados em algum lugar no início do seu aplicativo PyQt5
NUMERO_ATA_GLOBAL = None  # Deve ser definido em algum ponto antes de iniciar_processo
GERADOR_NUMERO_ATA = None  # Deve ser definido em algum ponto antes de iniciar_processo

def adicione_texto_formatado(paragraph, text, bold=False):
    run = paragraph.add_run(text)
    run.bold = bold
    font = run.font
    font.name = 'Calibri'
    font.size = Pt(12)

def seu_gerador_inicial(valor_inicial: int):
    """Gerador que fornece números de ata incrementais a partir de um número inicial."""
    numero = valor_inicial
    while True:
        valor_recebido = (yield numero)
        if valor_recebido is not None:
            numero = valor_recebido
        else:
            numero += 1

def confirmar_numero_ata(numero_ata, parent=None):
    global NUMERO_ATA_GLOBAL, GERADOR_NUMERO_ATA
    
    NUMERO_ATA_GLOBAL = int(numero_ata)
    mensagem = f'A próxima ata de registro de preços será "{numero_ata}-00"'
    QMessageBox.information(parent, "Confirmação", mensagem)
    
    GERADOR_NUMERO_ATA = seu_gerador_inicial(NUMERO_ATA_GLOBAL)
    next(GERADOR_NUMERO_ATA)

def iniciar_processo():
    if NUMERO_ATA_GLOBAL is None:
        raise ValueError("NUMERO_ATA not set!")
    
    criar_pastas_com_subpastas()
    processar_ata(NUMERO_ATA_GLOBAL)

def criar_pastas_com_subpastas(dataframe) -> None:

    if dataframe is None:
        QMessageBox.warning(None, "Erro", "Padrão de pregão não encontrado. Por favor, carregue um database antes de continuar.")
        print("Padrão de pregão não encontrado. Necessário carregar um database.")
        return
    
    relatorio_path = get_relatorio_path()
    combinacoes = dataframe[['num_pregao', 'ano_pregao', 'empresa']].drop_duplicates().values
    
    for num_pregao, ano_pregao, empresa in combinacoes:
        if pd.isna(num_pregao) or pd.isna(ano_pregao) or pd.isna(empresa):
            continue

        print(f"Original: {empresa}")  # Print antes da limpeza
        empresa_limpa = limpar_nome_empresa(empresa)
        print(f"Limpo: {empresa_limpa}")  # Print após a limpeza
        
        nome_dir_principal = f"PE {int(num_pregao)}-{int(ano_pregao)}"
        path_dir_principal = relatorio_path / Path(nome_dir_principal)
        
        if not path_dir_principal.exists():
            path_dir_principal.mkdir(parents=True)
            print(f"Criado diretório principal: {path_dir_principal}")
        
        if empresa not in NOMES_INVALIDOS and empresa:
            nome_subpasta = empresa_limpa
            path_subpasta = path_dir_principal / Path(nome_subpasta)
        
            if not path_subpasta.exists():
                path_subpasta.mkdir(parents=True)
                print(f"Criado subdiretório: {path_subpasta}")


def limpar_nome_empresa(nome_empresa):
    # Substituir caracteres não permitidos por "_" ou remover
    caracteres_invalidos = ['<', '>', ':', '"', '/', '\\', '|', '?', '*']
    for char in caracteres_invalidos:
        nome_empresa = nome_empresa.replace(char, '_')

    # Remover pontos extras apenas no final do nome da empresa
    if nome_empresa.endswith('.'):
        nome_empresa = nome_empresa.rstrip('.')

    return nome_empresa


import os
import subprocess
import sys

def abrir_pasta(pasta):
    if sys.platform == "win32":
        os.startfile(pasta)
    elif sys.platform == "darwin":  # MacOS
        subprocess.Popen(["open", pasta])
    else:  # Linux e outros
        subprocess.Popen(["xdg-open", pasta])

def processar_ata(NUMERO_ATA: int, nup_data, dataframe):
    relatorio_path = get_relatorio_path()

    if nup_data:
        nup = nup_data['nup']
    else:
        nup = "(INSIRA O NUP)"

    combinacoes = dataframe[['uasg', 'num_pregao', 'ano_pregao', 'empresa']].drop_duplicates().values
    NUMERO_ATA_atualizado = NUMERO_ATA

    for uasg, num_pregao, ano_pregao, empresa in combinacoes:
        if pd.isna(num_pregao) or pd.isna(ano_pregao) or pd.isna(empresa):
            continue

        try:
            uasg = int(uasg)
            num_pregao = int(num_pregao)
            ano_pregao = int(ano_pregao)
        except ValueError:
            # Mantenha os valores originais se a conversão falhar
            pass

        if empresa not in NOMES_INVALIDOS and empresa:
            nome_dir_principal = f"PE {int(num_pregao)}-{int(ano_pregao)}"
            path_dir_principal = relatorio_path / nome_dir_principal
            nome_subpasta = f"{empresa}"
            path_subpasta = path_dir_principal / nome_subpasta
            
            # Create subfolder if it doesn't exist
            if not path_subpasta.exists():
                path_subpasta.mkdir(parents=True, exist_ok=True)
            
            # Find the relevant record for this document
            registros_empresa = dataframe[dataframe['empresa'] == empresa]
            if not registros_empresa.empty:
                registro = registros_empresa.iloc[0].to_dict()
                itens_relacionados = registros_empresa.to_dict('records')
                email = registro.get("email", "E-mail não fornecido")  # Substitua 'E-mail não fornecido' se necessário

                # Construct the header text
                texto_substituto = f"Nº {uasg}/2024-{NUMERO_ATA_atualizado:03}/00\nPregão Eletrônico nº {num_pregao}/{ano_pregao}"
                num_contrato = f"Nº {uasg}/2024-{NUMERO_ATA_atualizado:03}/00"
                # Renderizar e salvar o documento
                tpl = DocxTemplate(TEMPLATE_PATH)

                itens_relacionados = registros_empresa.to_dict('records')

                soma_valor_homologado = gerar_soma_valor_homologado(itens_relacionados)

                context = {
                    "num_pregao": str(num_pregao),
                    "ano_pregao": str(ano_pregao),
                    "empresa": empresa,
                    "uasg": str(uasg),
                    "numero_ata": NUMERO_ATA_atualizado,
                    "soma_valor_homologado": soma_valor_homologado,
                    "cabecalho": texto_substituto,
                    "contrato": num_contrato,
                    "endereco": registro["endereco"],
                    "cnpj": registro["cnpj"],
                    "objeto": registro["objeto"],
                    "ordenador_despesa": registro["ordenador_despesa"],
                    "responsavel_legal": registro["responsavel_legal"],
                    "nup": nup,
                    "email": email 
                }
                
                tpl.render(context)
                nome_documento = f"{empresa} ata.docx"
                path_documento = path_subpasta / nome_documento

                try:
                    # Código para criar o arquivo TXT
                    nome_arquivo_txt = "E-mail.txt"
                    path_arquivo_txt = path_subpasta / nome_arquivo_txt
                    with open(path_arquivo_txt, "w") as arquivo_txt:
                        texto = (f"{email}\n\n"
                                f"Sr. Representante.\n\n"
                                f"Encaminho em anexo a Vossa Senhoria a ATA {num_contrato} "
                                f"decorrente do Pregão Eletrônico (SRP) nº {num_pregao}/{ano_pregao}, do Centro "
                                f"de Intendêcia da Marinha em Brasília (CeIMBra).\n\n"
                                f"Os documentos deverão ser conferidos, assinados e devolvidos a este Comando.\n\n"
                                f"A empresa receberá uma via, devidamente assinada, após a publicação.\n\n"
                                f"Respeitosamente,\n")
                        arquivo_txt.write(texto)

                    tpl.save(path_documento)
                    # Após salvar, modificar o documento com as informações adicionais
                    alterar_documento_criado(path_documento, registro, registro["cnpj"], itens_relacionados)
                    # Atualizar o NUMERO_ATA para o próximo valor
                    NUMERO_ATA_atualizado += 1
                except FileNotFoundError as e:
                    print(f"Erro ao salvar o documento: {e}")
            else:
                print(f"Nenhum registro encontrado para a empresa: {empresa}")

    # Atualizar a coluna num_ata após o loop para evitar o incremento prematuro
    if 'numero_da_ata' not in dataframe.columns:
        dataframe['numero_da_ata'] = ""
    dataframe['numero_da_ata'] = dataframe['numero_da_ata'].astype(str)
    for uasg, num_pregao, ano_pregao, empresa in combinacoes:
        dataframe.loc[dataframe['empresa'] == empresa, 'numero_da_ata'] = f"{uasg}/2023-{NUMERO_ATA_atualizado:03}/00"

    abrir_pasta(str(path_dir_principal))

    return NUMERO_ATA_atualizado  # Retornar o último número de ATA utilizado

def processar_contrato(NUMERO_ATA: int, nup_data):
    relatorio_path = get_relatorio_path()

    df = pd.read_csv(CSV_SICAF_PATH)
    df = df.dropna(subset=['num_pregao', 'ano_pregao', 'empresa'])

    if nup_data:
        nup = nup_data['nup']
    else:
        nup = "(INSIRA O NUP)"

    combinacoes = df[['uasg', 'num_pregao', 'ano_pregao', 'empresa']].drop_duplicates().values
    NUMERO_ATA_atualizado = NUMERO_ATA

    for uasg, num_pregao, ano_pregao, empresa in combinacoes:
        if pd.isna(num_pregao) or pd.isna(ano_pregao) or pd.isna(empresa):
            continue

        try:
            uasg = int(uasg)
            num_pregao = int(num_pregao)
            ano_pregao = int(ano_pregao)
        except ValueError:
            # Mantenha os valores originais se a conversão falhar
            pass

        if empresa not in NOMES_INVALIDOS and empresa:
            nome_dir_principal = f"PE {int(num_pregao)}-{int(ano_pregao)}"
            path_dir_principal = relatorio_path / nome_dir_principal
            nome_subpasta = f"{empresa}"
            path_subpasta = path_dir_principal / nome_subpasta
            
            # Create subfolder if it doesn't exist
            if not path_subpasta.exists():
                path_subpasta.mkdir(parents=True, exist_ok=True)
            
            # Find the relevant record for this document
            registros_empresa = df[df['empresa'] == empresa]
            if not registros_empresa.empty:
                registro = registros_empresa.iloc[0].to_dict()
                itens_relacionados = registros_empresa.to_dict('records')
                uasg_str = str(uasg)
                uasg_ultimos_5 = uasg_str[-5:]
                # Construct the header text
                texto_substituto = f"Nº {uasg_ultimos_5}/2023-{NUMERO_ATA_atualizado:03}/00\nPregão Eletrônico nº {num_pregao}/{ano_pregao}"
                num_contrato = f"Nº {uasg_ultimos_5}/2023-{NUMERO_ATA_atualizado:03}/00"
                # Renderizar e salvar o documento
                tpl = DocxTemplate(TEMPLATE_CONTRATO_PATH)

                itens_relacionados = registros_empresa.to_dict('records')

                soma_valor_homologado = gerar_soma_valor_homologado(itens_relacionados)

                context = {
                    "num_pregao": str(num_pregao),
                    "ano_pregao": str(ano_pregao),
                    "empresa": empresa,
                    "uasg": str(uasg),
                    "numero_ata": NUMERO_ATA_atualizado,
                    "soma_valor_homologado": soma_valor_homologado,
                    "cabecalho": texto_substituto,
                    "contrato": num_contrato,
                    "endereco": registro["endereco"],
                    "cnpj": registro["cnpj"],
                    "objeto": registro["objeto"],
                    "ordenador_despesa": registro["ordenador_despesa"],
                    "responsavel_legal": registro["responsavel_legal"],
                    "nup": nup 
                }
                tpl.render(context)
                nome_documento = f"{empresa} contrato.docx"
                path_documento = path_subpasta / nome_documento

                try:
                    tpl.save(path_documento)
                    # Após salvar, modificar o documento com as informações adicionais
                    alterar_contrato_criado(path_documento, registro, registro["cnpj"], itens_relacionados)
                    # Atualizar o NUMERO_ATA para o próximo valor
                    NUMERO_ATA_atualizado += 1
                except FileNotFoundError as e:
                    print(f"Erro ao salvar o documento: {e}")
            else:
                print(f"Nenhum registro encontrado para a empresa: {empresa}")

    # Atualizar a coluna num_ata após o loop para evitar o incremento prematuro
    if 'num_ata' not in df.columns:
        df['num_ata'] = ""
    df['num_ata'] = df['num_ata'].astype(str)
    for uasg, num_pregao, ano_pregao, empresa in combinacoes:
        df.loc[df['empresa'] == empresa, 'num_ata'] = f"{uasg}/2024-{NUMERO_ATA_atualizado:03}/00"

    # Salvar o DataFrame atualizado
    csv_filename = f"PE {int(num_pregao)}-{int(ano_pregao)}.csv"
    df.to_csv(csv_filename, index=False)
    excel_filename = f"PE {int(num_pregao)}-{int(ano_pregao)}.xlsx"
    df.to_excel(excel_filename, index=False)

    abrir_pasta(str(path_dir_principal))

    return NUMERO_ATA_atualizado  # Retornar o último número de ATA utilizado

def gerar_soma_valor_homologado(itens):
    valor_total = sum(float(item["valor_homologado_total_item"] or 0) for item in itens)  # tratando None como 0
    valor_extenso = valor_por_extenso(valor_total)
    return f'R$ {formatar_brl(valor_total)} ({valor_extenso})'

def alterar_documento_criado(caminho_documento, registro, cnpj, itens):
    # Carregando o documento real
    doc = Document(caminho_documento)
    
    # Iterando por cada parágrafo do documento
    for paragraph in doc.paragraphs:
        if '{relacao_empresa}' in paragraph.text:
            # Substituindo o marcador pelo conteúdo gerado pela função inserir_relacao_empresa
            paragraph.clear()  # Limpar o parágrafo atual
            inserir_relacao_empresa(paragraph, registro, cnpj)
        
        # Verificando o marcador {relacao_item}
        if '{relacao_item}' in paragraph.text:
            # Substituindo o marcador pelo conteúdo gerado pela função inserir_relacao_itens
            paragraph.clear()  # Limpar o parágrafo atual
            inserir_relacao_itens(paragraph, itens)
    
    # Salvando as alterações no documento
    doc.save(caminho_documento)

def inserir_relacao_empresa(paragrafo, registro, cnpj):
    dados = {
        "Razão Social": registro["empresa"],
        "CNPJ": registro["cnpj"],
        "Endereço": registro["endereco"],
        "Município-UF": registro["municipio"],
        "CEP": registro["cep"],
        "Telefone": registro["telefone"],
        "E-mail": registro["email"]
    }

    total_itens = len(dados)
    contador = 1
    
    for chave, valor in dados.items():
        adicione_texto_formatado(paragrafo, f"{chave}: ", True)

        # Verifica se é a penúltima linha
        if contador == total_itens - 1:
            adicione_texto_formatado(paragrafo, f"{valor}; e\n", False)
        # Verifica se é a última linha
        elif contador == total_itens:
            adicione_texto_formatado(paragrafo, f"{valor}.\n", False)
        else:
            adicione_texto_formatado(paragrafo, f"{valor};\n", False)

        contador += 1
    
    adicione_texto_formatado(paragrafo, "Representada neste ato, por seu representante legal, o(a) Sr(a) ", False)
    adicione_texto_formatado(paragrafo, f'{registro["responsavel_legal"]}.\n', False)

def gerar_campos_item(item):
    # Convertendo item_num e catalogo para inteiros
    item_num_int = int(item["item_num"])
    catalogo_int = int(item["catalogo"])

    # Formatando a quantidade
    quantidade_formatada = f"{float(item['quantidade']):.2f}".rstrip('0').rstrip('.')
    descricao_detalhada_ajustada = item["descricao_detalhada"].replace("\n", " ")

    #Padrão Material
    return [
        (f'Item {item_num_int} - {item["descricao_tr"]} | Catálogo: {catalogo_int}', True),
        (f'Descrição: {descricao_detalhada_ajustada}', False),
        (f'Unidade de Fornecimento: {item["unidade"]}', False),
        (f'Marca/Fabricante: {item["marca_fabricante"]}   |   Modelo/Versão: {item["modelo_versao"]}', False),
        (f'Quantidade: {quantidade_formatada}   |   Valor Unitário: R$ {formatar_brl(item["valor_homologado_item_unitario"])}   |   Valor Total do Item: R$ {formatar_brl(item["valor_homologado_total_item"])}', False),
        (f'{"-" * 130}', False)
    ]
    #Padrão Serviços
    # return [
    #     (f'Item {item_num_int}', False),
    #     (f'Descrição: {item["descricao_detalhada_tr"]}', False),
    #     (f'Unidade de Fornecimento: {item["unidade"]}', False),
    #     # (f'Marca/Fabricante: {item["marca_fabricante"]}   |   Modelo/Versão: {item["modelo_versao"]}', False),
    #     (f'Quantidade: {quantidade_formatada}   |   Valor Unitário: R$ {formatar_brl(item["valor_homologado_item_unitario"])}   |   Valor Total do Item: R$ {formatar_brl(item["valor_homologado_total_item"])}', False),
    #     (f'{"-" * 130}', False)
    # ]
    # Padrão Material

def inserir_relacao_itens(paragrafo, itens):
    # Primeiro, limpamos o parágrafo para remover o placeholder e qualquer outro texto.
    paragrafo.clear()

    for item in itens:
        campos = gerar_campos_item(item)
        for texto, negrito in campos:
            adicione_texto_formatado(paragrafo, texto + '\n', negrito)

    # Calculando o valor total homologado para a empresa
    valor_total = sum(float(item["valor_homologado_total_item"] or 0) for item in itens)  # tratando None como 0
    valor_extenso = valor_por_extenso(valor_total)
    
    # Criando a string formatada
    texto_soma_valor_homologado = f'R$ {formatar_brl(valor_total)} ({valor_extenso})'

    # Inserindo o texto formatado no parágrafo
    adicione_texto_formatado(paragrafo, 'Valor total homologado para a empresa:\n', False)
    adicione_texto_formatado(paragrafo, texto_soma_valor_homologado + '\n', True)

    # Retornando o texto formatado
    return texto_soma_valor_homologado

def formatar_brl(valor):
    try:
        if valor is None:
            return "Não disponível"  # Retorna uma string informativa caso o valor seja None
        # Formata o número no formato monetário brasileiro sem utilizar a biblioteca locale
        return f"R$ {float(valor):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except (ValueError, TypeError):
        return "Valor inválido"  # Retorna isso se não puder converter para float
       
def valor_por_extenso(valor):
    extenso = num2words(valor, lang='pt_BR', to='currency')
    return extenso.capitalize()


def alterar_contrato_criado(caminho_documento, registro, cnpj, itens):
    # Carregando o documento real
    doc = Document(caminho_documento)
    
    # Iterando por cada parágrafo do documento
    for paragraph in doc.paragraphs:
        if '{relacao_empresa}' in paragraph.text:
            # Substituindo o marcador pelo conteúdo gerado pela função inserir_relacao_empresa
            paragraph.clear()  # Limpar o parágrafo atual
            inserir_relacao_empresa(paragraph, registro, cnpj)
        
        # Verificando o marcador {relacao_item}
        if '{relacao_item}' in paragraph.text:
            # Substituindo o marcador pelo conteúdo gerado pela função inserir_relacao_itens
            paragraph.clear()  # Limpar o parágrafo atual
            inserir_relacao_itens(paragraph, itens)
    
    # Salvando as alterações no documento
    doc.save(caminho_documento)

# def inserir_relacao_itens(paragrafo, itens):
#     # Primeiro, limpamos o parágrafo para remover o placeholder e qualquer outro texto.
#     paragrafo.clear()

#     for item in itens:
#         campos = gerar_campos_item(item)
#         for texto, negrito in campos:
#             adicione_texto_formatado(paragrafo, texto + '\n', negrito)
